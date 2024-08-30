# -*- coding: utf-8 -*-
import sqlite3
import base64
import requests
import json
import hashlib
import itertools
import logging
from multiprocessing.util import register_after_fork
import os
import base64
import hashlib
import pprint
import re
import shutil
import sndhdr
import tempfile
import transaction
import zipfile

import collections
from collections import defaultdict

from pathvalidate import sanitize_filename

import urllib
from urllib import request

import minio

from pydub import AudioSegment

import sqlalchemy
from sqlalchemy.orm.exc import NoResultFound
from sqlalchemy import and_, create_engine, func, literal

from lingvodoc.models import (
    Client,
    DBSession as SyncDBSession,
    UserBlobs,
    TranslationAtom,
    TranslationGist,
    Field,
    Entity,
    LexicalEntry,
    Dictionary,
    Language,
    User,
    DictionaryPerspectiveToField,
    DictionaryPerspective,
    BaseGroup,
    Group,
    PublishingEntity
)

from lingvodoc.cache.caching import TaskStatus, initialize_cache
from lingvodoc.utils import explain_analyze, sanitize_worksheet_name
from lingvodoc.utils.search import recursive_sort, translation_gist_id_search
from lingvodoc.utils.static_fields import fields_static
from lingvodoc.views.v2.utils import as_storage_file, storage_file

from sqlalchemy.orm import (
    aliased, sessionmaker,
)

from sqlalchemy import tuple_, and_, or_
from lingvodoc.scripts import elan_parser
from pdb import set_trace as A
import io, codecs
from xlsxwriter import Workbook as xlsxDocument
from docx import Document as docxDocument
from docx.shared import Inches

from PyRTF.Elements import Document as rtfDocument
from PyRTF.document.section import Section
from PyRTF.document.paragraph import Cell, Paragraph, Table
from PyRTF.PropertySets import BorderPropertySet, FramePropertySet, TabPropertySet
import rtfunicode

# from time import time
import datetime
import time
import traceback
from os import path, makedirs
from errno import EEXIST
from shutil import copyfileobj
import uuid


log = logging.getLogger(__name__)

EAF_TIERS = {
    "literary translation": "Translation of Paradigmatic forms",
    "text": "Transcription of Paradigmatic forms",
    "word": "Word",
    "transcription": "Transcription",
    "translation": "Translation"
}


def find_lexical_entries_by_tags(tags, session, published):
    result = session.query(LexicalEntry) \
        .join(Entity) \
        .join(PublishingEntity) \
        .filter(
            LexicalEntry.marked_for_deletion == False,
            Entity.content.in_(tags),
            Entity.marked_for_deletion == False,
            Entity.field_client_id == 66,
            Entity.field_object_id == 25,
            PublishingEntity.accepted == True)
    if published is not None:
        result = result.filter(PublishingEntity.published == published)
    return result.all()


def find_all_tags(tag, session, published):
    tags = {tag}
    new_tags = {tag}
    while new_tags:
        lexical_entries = find_lexical_entries_by_tags(new_tags, session, published)
        new_tags = set()
        for lex in lexical_entries:
            entities = session.query(Entity)  \
                .join(PublishingEntity) \
                .filter(
                    Entity.marked_for_deletion == False,
                    Entity.parent == lex,
                    Entity.field_client_id == 66,
                    Entity.field_object_id == 25,
                    PublishingEntity.accepted == True)
            if published is not None:
                entities = entities.filter(PublishingEntity.published == published)
            for entity in entities:
                if entity.content not in tags:
                    tags.add(entity.content)
                    new_tags.add(entity.content)
    return tags


def find_group_by_tags(
    session,
    entry_id_set,
    tag_set,
    tag_list,
    field_client_id,
    field_object_id,
    published = None):
    """
    Optimized retrieval of a group of lexical entries linked by a set of tags from a specified field.
    """

    # While we have tags we don't have all lexical entries for,
    # we get these all entries of these tags...

    while tag_list:

        entry_id_query = (

            session.query(
                LexicalEntry.client_id,
                LexicalEntry.object_id)

            .filter(
                LexicalEntry.marked_for_deletion == False,
                Entity.parent_id == LexicalEntry.id,
                Entity.field_id == (field_client_id, field_object_id),
                Entity.marked_for_deletion == False,
                Entity.content.in_(tag_list),
                PublishingEntity.id == Entity.id,
                PublishingEntity.accepted == True))

        if published is not None:

            entry_id_query = entry_id_query.filter(
                PublishingEntity.published == published)

        entry_id_list = []

        for entry_id in entry_id_query.distinct():

            if entry_id not in entry_id_set:

                entry_id_set.add(entry_id)
                entry_id_list.append(entry_id)

        if not entry_id_list:
            break

        # And then get all tags for entries we haven't already done it for.

        tag_query = (

            session.query(
                Entity.content)

            .filter(
                tuple_(Entity.parent_client_id, Entity.parent_object_id)
                    .in_(entry_id_list),
                Entity.field_id == (field_client_id, field_object_id),
                Entity.marked_for_deletion == False,
                PublishingEntity.id == Entity.id,
                PublishingEntity.accepted == True))

        if published is not None:

            tag_query = tag_query.filter(
                PublishingEntity.published == published)

        tag_list = []

        for tag in tag_query.distinct():

            if tag not in tag_set:

                tag_set.add(tag)
                tag_list.append(tag)

    # Showing what we've found.

    log.debug(
        '\nfind_group_by_tags({0}, {1}):'
        '\nlen(tag_set): {2}'
        '\nlen(entry_id_set): {3}'.format(
        field_client_id,
        field_object_id,
        len(tag_set),
        len(entry_id_set)))

    return entry_id_set, tag_set


def generate_massive_cell_simple(
    tag,
    session,
    text_fields,
    published):
    """
    Gathers info of etymologically linked lexical entries via standard simple method.
    """

    result_list = []

    # Current standard method.

    tags = find_all_tags(tag, session, published)
    lexical_entries = find_lexical_entries_by_tags(tags, session, published)

    for lex in lexical_entries:
        entities = session.query(Entity).join(PublishingEntity) \
            .filter(Entity.parent_id == (lex.client_id, lex.object_id),
                    PublishingEntity.accepted == True,
                    Entity.marked_for_deletion == False)
        if published is not None:
            entities = entities.filter(PublishingEntity.published == published)
        subres = []
        for entity in entities:
            if (entity.field_client_id, entity.field_object_id) in text_fields and entity.content is not None:
                subres.append(entity.content)
        if len(subres) > 0:
            result_list.append("; ".join(subres))

    return "\n".join(result_list)


def generate_massive_cell_optimized(
    tag,
    session,
    text_fields,
    published):
    """
    Optimized gathering of etymologically linked lexical entries info.
    """

    entry_id_set, tag_set = (

        find_group_by_tags(
            session, set(), {tag}, [tag], 66, 25, published))

    query = (
            
        session

            .query()
            
            .filter(
                Entity.marked_for_deletion == False,
                tuple_(Entity.parent_client_id, Entity.parent_object_id)
                    .in_(entry_id_set),
                tuple_(Entity.field_client_id, Entity.field_object_id)
                    .in_(text_fields),
                Entity.content != None,
                PublishingEntity.id == Entity.id,
                PublishingEntity.accepted == True)

            .add_columns(
                func.string_agg(Entity.content, '; ')
                .label('content'))
            
            .group_by(
                Entity.parent_client_id, Entity.parent_object_id))

    if published is not None:

        query = query.filter(
            PublishingEntity.published == published)

    return '\n'.join(
        row[0] for row in query)


def generate_massive_cell_cte(
    tag,
    session,
    text_fields,
    published):
    """
    Gathering of etymologically linked lexical entries info using recursive CTE.

    Takes much more time then the back-and-forth implementation in generate_massive_cell_optimized() due to
    inefficiences stemming from the recursive CTE query restrictions.
    """

    entry_id_query = (
            
        session.query(
            LexicalEntry.client_id,
            LexicalEntry.object_id)

        .filter(
            LexicalEntry.marked_for_deletion == False,
            Entity.parent_id == LexicalEntry.id,
            Entity.field_client_id == 66,
            Entity.field_object_id == 25,
            Entity.marked_for_deletion == False,
            Entity.content == tag,
            PublishingEntity.id == Entity.id,
            PublishingEntity.accepted == True))

    if published is not None:

        entry_id_query = entry_id_query.filter(
            PublishingEntity.published == published)

    # See https://stackoverflow.com/questions/53186429/sqlalchemy-simple-recursive-cte-query for Sqlalchemy
    # resursive CTE example.

    cte_query = entry_id_query.cte(
        recursive = True, name = 'entry_id')

    E_entry = aliased(Entity, name = 'E1')
    E_tag = aliased(Entity, name = 'E2')

    P_entry = aliased(PublishingEntity, name = 'P1')
    P_tag = aliased(PublishingEntity, name = 'P2')

    recursive_query = (

        session.query(
            LexicalEntry.client_id,
            LexicalEntry.object_id)
        
        .filter(
            LexicalEntry.marked_for_deletion == False,
            E_entry.parent_id == LexicalEntry.id,
            E_entry.field_client_id == 66,
            E_entry.field_object_id == 25,
            E_entry.marked_for_deletion == False,
            P_entry.id == E_entry.id,
            P_entry.accepted == True,
            E_entry.content == E_tag.content,
            E_tag.parent_id == (cte_query.c.client_id, cte_query.c.object_id),
            E_tag.field_client_id == 66,
            E_tag.field_object_id == 25,
            E_tag.marked_for_deletion == False,
            P_tag.id == E_tag.id,
            P_tag.accepted == True))

    if published is not None:

        recursive_query = (
                
            recursive_query.filter(
                P_entry.published == published,
                P_tag.published == published))

    cte_query = cte_query.union(
        recursive_query.distinct())

    # Query for required cell data.

    content_query = (
            
        session

            .query()
            
            .filter(
                Entity.marked_for_deletion == False,
                tuple_(Entity.parent_client_id, Entity.parent_object_id)
                    .in_(session.query(cte_query)),
                tuple_(Entity.field_client_id, Entity.field_object_id)
                    .in_(text_fields),
                Entity.content != None,
                PublishingEntity.id == Entity.id,
                PublishingEntity.accepted == True)

            .add_columns(
                func.string_agg(Entity.content, '; ')
                .label('content'))
            
            .group_by(
                Entity.parent_client_id, Entity.parent_object_id))

    if published is not None:

        content_query = content_query.filter(
            PublishingEntity.published == published)

    log.debug(
        '\ncontent_query:\n' +
        str(content_query))

    return '\n'.join(
        row[0] for row in content_query)


def generate_massive_cell_temp_table(
    tag,
    session,
    text_field_id_table_name,
    published,
    __debug_flag__ = False):
    """
    Gathering of etymologically linked lexical entries info using temporary tables.
    """

    uuid_str = str(uuid.uuid4()).replace('-', '_')

    entry_id_table_name = 'entry_id_table_' + uuid_str
    tag_table_name = 'tag_table_' + uuid_str

    tag_list_name = 'tag_list_' + uuid_str
    tag_list_prev_name = 'tag_list_prev_' + uuid_str

    published_str = (
        '' if published is None else
        ' and P.published = ' + str(published).lower())

    # Temporary table for lexical entry ids.

    sql_str = ('''

        create temporary table
        
        {entry_id_table_name} (
          client_id BIGINT,
          object_id BIGINT,
          primary key (client_id, object_id))

        on commit drop;

        insert into {entry_id_table_name}

        select
          L.client_id,
          L.object_id

        from
          lexicalentry L,
          public.entity E,
          publishingentity P

        where
          L.marked_for_deletion = false and
          E.parent_client_id = L.client_id and
          E.parent_object_id = L.object_id and
          E.field_client_id = 66 and
          E.field_object_id = 25 and
          E.marked_for_deletion = false and
          E.content = :tag and
          P.client_id = E.client_id and
          P.object_id = E.object_id and
          P.accepted = true{0}

        on conflict do nothing;

        '''.format(
            published_str,
            entry_id_table_name = entry_id_table_name))

    session.execute(
        sql_str, {'tag': tag})

    # Temporary table for tags.

    sql_str = ('''

        create temporary table

        {tag_table_name} (
          tag TEXT primary key)

        on commit drop;

        insert into {tag_table_name}
          values (:tag);

        create temporary table
        
        {tag_list_name}
        
        on commit drop as
        with

        tag_cte as (

          insert into {tag_table_name}
          select distinct E.content

          from
            public.entity E,
            publishingentity P

          where
            (E.parent_client_id, E.parent_object_id) in (
              select * from {entry_id_table_name}) and
            E.field_client_id = 66 and
            E.field_object_id = 25 and
            E.marked_for_deletion = false and
            P.client_id = E.client_id and
            P.object_id = E.object_id and
            P.accepted = true{0}

          on conflict do nothing
          returning *)
        
        select * from tag_cte;

        '''.format(
            published_str,
            entry_id_table_name = entry_id_table_name,
            tag_table_name = tag_table_name,
            tag_list_name = tag_list_name))

    session.execute(
        sql_str, {'tag': tag})

    # While we have tags we don't have all lexical entries for, we get these all new entries of these tags
    # and then all new tags of this new entries.

    exists_str = '''

        select exists (
          select 1 from {tag_list_name});
          
        '''.format(
            tag_list_name = tag_list_name)

    while session.execute(exists_str).scalar():

        sql_str = ('''

            alter table {tag_list_name}
              rename to {tag_list_prev_name};

            create temporary table
            
            {tag_list_name}
            
            on commit drop as
            with

            entry_id_cte as (

              insert into {entry_id_table_name}

              select distinct
                L.client_id,
                L.object_id

              from
                lexicalentry L,
                public.entity E,
                publishingentity P

              where
                L.marked_for_deletion = false and
                E.parent_client_id = L.client_id and
                E.parent_object_id = L.object_id and
                E.field_client_id = 66 and
                E.field_object_id = 25 and
                E.marked_for_deletion = false and
                E.content in (
                  select * from {tag_list_prev_name}) and
                P.client_id = E.client_id and
                P.object_id = E.object_id and
                P.accepted = true{0}

              on conflict do nothing
              returning *),

            tag_cte as (

              insert into {tag_table_name}
              select distinct E.content

              from
                public.entity E,
                publishingentity P

              where
                (E.parent_client_id, E.parent_object_id) in (
                  select * from entry_id_cte) and
                E.field_client_id = 66 and
                E.field_object_id = 25 and
                E.marked_for_deletion = false and
                P.client_id = E.client_id and
                P.object_id = E.object_id and
                P.accepted = true{0}

              on conflict do nothing
              returning *)
            
            select * from tag_cte;

            drop table {tag_list_prev_name};

            '''.format(
                published_str,
                entry_id_table_name = entry_id_table_name,
                tag_table_name = tag_table_name,
                tag_list_name = tag_list_name,
                tag_list_prev_name = tag_list_prev_name))

        session.execute(
            sql_str)

    # And now we get required info of the linked lexical entries.

    query = (
            
        session

            .query()
            
            .filter(
                Entity.marked_for_deletion == False,

                tuple_(
                    Entity.parent_client_id, Entity.parent_object_id)

                    .in_(sqlalchemy.text(
                        'select * from ' + entry_id_table_name)),

                tuple_(
                    Entity.field_client_id, Entity.field_object_id)

                    .in_(sqlalchemy.text(
                        'select * from ' + text_field_id_table_name)),

                Entity.content != None,
                PublishingEntity.id == Entity.id,
                PublishingEntity.accepted == True)

            .add_columns(
                func.string_agg(Entity.content, '; ')
                .label('content'))
            
            .group_by(
                Entity.parent_client_id, Entity.parent_object_id))

    if published is not None:

        query = query.filter(
            PublishingEntity.published == published)

    if __debug_flag__:

        # Showing PostgreSQL's explain analyze, if required.

        row_list = (
            self.session.execute(
                explain_analyze(query)).fetchall())

        log.debug(''.join(
            '\n' + row[0] for row in row_list))

    return '\n'.join(
        row[0] for row in query)


def generate_massive_cell_plpgsql(
    tag,
    session,
    published):
    """
    Gathering of etymologically linked lexical entries info using stored PL/pgSQL procedure.
    """

    row_list = (
            
        session.execute(
            'select etymology_text(:tag, :publish)',
            {'tag': tag, 'publish': published}).fetchall())

    return '\n'.join(
        row[0] for row in row_list)


class Save_Context(object):
    """
    Holds data of the saving of dictionary/dictionaries into an XLSX workbook.

    Also used in advanced_search.
    """

    def __init__(
        self,
        locale_id,
        session,
        cognates_flag = True,
        sound_flag = False,
        markup_flag = False,
        storage = None,
        f_type = 'xlsx',
        __debug_flag__ = False):

        self.locale_id = locale_id
        self.session = session

        self.cognates_flag = (cognates_flag and
                              f_type == 'xlsx')
        self.sound_flag = sound_flag
        self.markup_flag = markup_flag

        self.stream = io.BytesIO()
        self.workbook = xlsxDocument(self.stream, {'in_memory': True}) if f_type == 'xlsx' else None
        self.document = docxDocument() if f_type == 'docx' else None
        self.richtext = rtfDocument() if f_type == 'rtf' else None

        # Getting up-to-date text field info.

        self.session.execute(
            'refresh materialized view text_field_id_view;')

        # Preparation for etymology listing.

        if cognates_flag:

            self.etymology_dict = {}
            self.perspective_info_dict = {}

            self.transcription_fid = fields_static['Transcription']
            self.translation_fid = fields_static['Translation']

            # We would need standard language ordering for ordering cognates, see resolve_language_tree().

            language_list = (
                    
                self.session
                    .query(Language)

                    .filter_by(
                        marked_for_deletion = False)

                    .order_by(
                        Language.parent_client_id,
                        Language.parent_object_id,
                        Language.additional_metadata['younger_siblings'])

                    .all())

            order_list = recursive_sort(language_list)

            self.language_order_dict = {
                (cid, oid): index
                for index, (level, cid, oid, language) in enumerate(order_list)}

            log.debug(
                '\nlanguage_order_dict:\n' +
                pprint.pformat(
                    self.language_order_dict, width = 192))

            # And we would need a mapping from perspectives to transcription / translation fields, and a
            # temporary table for cognate entries.

            uuid_str = str(uuid.uuid4()).replace('-', '_')

            self.pid_fid_table_name = 'pid_fid_table_' + uuid_str
            self.eid_pid_table_name = 'eid_pid_table_' + uuid_str

            sound_fid_str = (

                '''
                sound_client_id BIGINT,
                sound_object_id BIGINT,
                ''')

            markup_fid_str = (

                '''
                markup_client_id BIGINT,
                markup_object_id BIGINT,
                ''')

            self.session.execute('''

                create temporary table
                
                {pid_fid_table_name} (
                  perspective_client_id BIGINT,
                  perspective_object_id BIGINT,
                  transcription_client_id BIGINT,
                  transcription_object_id BIGINT,
                  translation_client_id BIGINT,
                  translation_object_id BIGINT,
                  third_field_client_id BIGINT,
                  third_field_object_id BIGINT,{sound_str}{markup_str}

                  primary key (
                    perspective_client_id,
                    perspective_object_id))

                on commit drop;

                create temporary table
                
                {eid_pid_table_name} (
                  entry_client_id BIGINT,
                  entry_object_id BIGINT,
                  perspective_client_id BIGINT,
                  perspective_object_id BIGINT,

                  primary key (
                    entry_client_id,
                    entry_object_id))

                on commit drop;

                '''.format(

                    pid_fid_table_name =
                        self.pid_fid_table_name,

                    eid_pid_table_name =
                        self.eid_pid_table_name,

                    sound_str =
                        sound_fid_str if sound_flag else '',

                    markup_str =
                        markup_fid_str if markup_flag else ''))

            # Some commonly used etymological SQL queries.

            self.sql_etymology_str_a = ('''

                truncate table {eid_pid_table_name};

                insert into
                {eid_pid_table_name}

                select
                L.client_id,
                L.object_id,
                L.parent_client_id,
                L.parent_object_id

                from
                lexicalentry L

                where
                (L.client_id, L.object_id) in
                  (select * from linked_group(66, 25, :tag, (:publish :: BOOLEAN)));

                select distinct
                perspective_client_id,
                perspective_object_id

                from
                {eid_pid_table_name};

                '''.format(
                    eid_pid_table_name = self.eid_pid_table_name))

            sound_fid_str = (

                ''',
                Tp.sound_client_id,
                Tp.sound_object_id
                ''')

            markup_fid_str = (

                ''',
                Tp.markup_client_id,
                Tp.markup_object_id
                ''')

            sound_cte_str = (
                    
                ''',

                sound_cte as (

                  select
                  E.parent_client_id,
                  E.parent_object_id,
                  jsonb_agg(
                    jsonb_build_array(
                      E.content,
                      extract(epoch from E.created_at))) content_list

                  from
                  eid_pid_fid_cte T,
                  public.entity E,
                  publishingentity P
                  
                  where
                  E.parent_client_id = T.entry_client_id and
                  E.parent_object_id = T.entry_object_id and
                  E.field_client_id = T.sound_client_id and
                  E.field_object_id = T.sound_object_id and
                  E.content is not null and
                  E.marked_for_deletion is false and
                  P.client_id = E.client_id and
                  P.object_id = E.object_id and
                  P.accepted = true{0}
                  
                  group by
                  E.parent_client_id,
                  E.parent_object_id)

                ''')

            sound_select_str = (

                ''',
                S.content_list
                ''')

            sound_join_str = (

                '''
                left outer join
                sound_cte S
                on
                S.parent_client_id = T.entry_client_id and
                S.parent_object_id = T.entry_object_id
                ''')

            markup_cte_str = (
                    
                ''',

                markup_cte as (

                  select
                  E.parent_client_id,
                  E.parent_object_id,
                  jsonb_agg(
                    jsonb_build_array(
                      E.content,
                      extract(epoch from E.created_at))) content_list

                  from
                  eid_pid_fid_cte T,
                  public.entity E,
                  publishingentity P
                  
                  where
                  E.parent_client_id = T.entry_client_id and
                  E.parent_object_id = T.entry_object_id and
                  E.field_client_id = T.markup_client_id and
                  E.field_object_id = T.markup_object_id and
                  E.content is not null and
                  E.marked_for_deletion is false and
                  P.client_id = E.client_id and
                  P.object_id = E.object_id and
                  P.accepted = true{0}
                  
                  group by
                  E.parent_client_id,
                  E.parent_object_id)

                ''')

            markup_select_str = (

                ''',
                M.content_list
                ''')

            markup_join_str = (

                '''
                left outer join
                markup_cte M
                on
                M.parent_client_id = T.entry_client_id and
                M.parent_object_id = T.entry_object_id
                ''')

            sound_markup_cte_str = (
                    
                ''',

                sound_markup_cte_a as (

                  select
                  Es.parent_client_id,
                  Es.parent_object_id,
                  Es.client_id,
                  Es.object_id,
                  Es.content,
                  Es.created_at,
                  jsonb_agg(
                    jsonb_build_array(
                      Em.content,
                      extract(epoch from Em.created_at))) markup_list

                  from
                  eid_pid_fid_cte T

                  inner join
                  public.entity Es

                  on
                  Es.parent_client_id = T.entry_client_id and
                  Es.parent_object_id = T.entry_object_id and
                  Es.field_client_id = T.sound_client_id and
                  Es.field_object_id = T.sound_object_id and
                  Es.content is not null and
                  Es.marked_for_deletion is false

                  inner join
                  publishingentity Ps

                  on
                  Ps.client_id = Es.client_id and
                  Ps.object_id = Es.object_id and
                  Ps.accepted = true{1}

                  left outer join public.entity Em

                  on
                  Em.parent_client_id = T.entry_client_id and
                  Em.parent_object_id = T.entry_object_id and
                  Em.field_client_id = T.markup_client_id and
                  Em.field_object_id = T.markup_object_id and
                  Em.content is not null and
                  Em.marked_for_deletion is false and
                  Em.self_client_id = Es.client_id and
                  Em.self_object_id = Es.object_id

                  left outer join publishingentity Pm

                  on
                  Pm.client_id = Em.client_id and
                  Pm.object_id = Em.object_id and
                  Pm.accepted = true{2}
                  
                  group by
                  Es.parent_client_id,
                  Es.parent_object_id,
                  Es.client_id,
                  Es.object_id),

                sound_markup_cte_b as (

                  select
                  parent_client_id,
                  parent_object_id,
                  jsonb_agg(
                    jsonb_build_array(
                      content,
                      extract(epoch from created_at),
                      markup_list)) content_list

                  from
                  sound_markup_cte_a
                  
                  group by
                  parent_client_id,
                  parent_object_id)

                ''')

            sound_markup_select_str = (

                ''',
                SM.content_list
                ''')

            sound_markup_join_str = (

                '''
                left outer join
                sound_markup_cte_b SM
                on
                SM.parent_client_id = T.entry_client_id and
                SM.parent_object_id = T.entry_object_id
                ''')

            sound_only_flag = sound_flag and not markup_flag
            markup_only_flag = markup_flag and not sound_flag
            sound_markup_flag = sound_flag and markup_flag

            self.sql_etymology_str_b = ('''

                with

                eid_pid_fid_cte as (

                  select
                  Te.entry_client_id,
                  Te.entry_object_id,
                  Te.perspective_client_id,
                  Te.perspective_object_id,
                  Tp.transcription_client_id,
                  Tp.transcription_object_id,
                  Tp.translation_client_id,
                  Tp.translation_object_id,
                  Tp.third_field_client_id,
                  Tp.third_field_object_id{sound_fid_str}{markup_fid_str}

                  from
                  {eid_pid_table_name} Te,
                  {pid_fid_table_name} Tp

                  where
                  Te.perspective_client_id = Tp.perspective_client_id and
                  Te.perspective_object_id = Tp.perspective_object_id),

                transcription_cte as (

                  select
                  E.parent_client_id,
                  E.parent_object_id,
                  array_agg(E.content) content_list

                  from
                  eid_pid_fid_cte T,
                  public.entity E,
                  publishingentity P
                  
                  where
                  E.parent_client_id = T.entry_client_id and
                  E.parent_object_id = T.entry_object_id and
                  E.field_client_id = T.transcription_client_id and
                  E.field_object_id = T.transcription_object_id and
                  E.content is not null and
                  E.marked_for_deletion = false and
                  P.client_id = E.client_id and
                  P.object_id = E.object_id and
                  P.accepted = true{{0}}
                  
                  group by
                  E.parent_client_id,
                  E.parent_object_id),

                translation_cte as (

                  select
                  E.parent_client_id,
                  E.parent_object_id,
                  array_agg(E.content) content_list

                  from
                  eid_pid_fid_cte T,
                  public.entity E,
                  publishingentity P
                  
                  where
                  E.parent_client_id = T.entry_client_id and
                  E.parent_object_id = T.entry_object_id and
                  E.field_client_id = T.translation_client_id and
                  E.field_object_id = T.translation_object_id and
                  E.content is not null and
                  E.marked_for_deletion = false and
                  P.client_id = E.client_id and
                  P.object_id = E.object_id and
                  P.accepted = true{{0}}
                  
                  group by
                  E.parent_client_id,
                  E.parent_object_id),

                third_field_cte as (

                  select
                  E.parent_client_id,
                  E.parent_object_id,
                  array_agg(E.content) content_list

                  from
                  eid_pid_fid_cte T,
                  public.entity E,
                  publishingentity P
                  
                  where
                  E.parent_client_id = T.entry_client_id and
                  E.parent_object_id = T.entry_object_id and
                  E.field_client_id = T.third_field_client_id and
                  E.field_object_id = T.third_field_object_id and
                  E.content is not null and
                  E.marked_for_deletion = false and
                  P.client_id = E.client_id and
                  P.object_id = E.object_id and
                  P.accepted = true{{0}}
                  
                  group by
                  E.parent_client_id,
                  E.parent_object_id){sound_markup_cte_str}

                select
                T.entry_client_id,
                T.entry_object_id,
                T.perspective_client_id,
                T.perspective_object_id,
                Xc.content_list,
                Xl.content_list,
                Tf.content_list{sound_markup_select_str}

                from
                eid_pid_fid_cte T

                left outer join
                transcription_cte Xc
                on
                Xc.parent_client_id = T.entry_client_id and
                Xc.parent_object_id = T.entry_object_id

                left outer join
                translation_cte Xl
                on
                Xl.parent_client_id = T.entry_client_id and
                Xl.parent_object_id = T.entry_object_id
                
                left outer join
                third_field_cte Tf
                on
                Tf.parent_client_id = T.entry_client_id and
                Tf.parent_object_id = T.entry_object_id{sound_markup_join_str};

                '''.format(

                    pid_fid_table_name =
                        self.pid_fid_table_name,

                    eid_pid_table_name =
                        self.eid_pid_table_name,
                        
                    sound_fid_str =
                        sound_fid_str if sound_flag else '',
                        
                    markup_fid_str =
                        markup_fid_str if markup_flag else '',
                        
                    sound_markup_cte_str =
                        sound_cte_str if sound_only_flag else
                        markup_cte_str if markup_only_flag else
                        sound_markup_cte_str if sound_markup_flag else '',
                        
                    sound_markup_select_str =
                        sound_select_str if sound_only_flag else
                        markup_select_str if markup_only_flag else
                        sound_markup_select_str if sound_markup_flag else '',

                    sound_markup_join_str =
                        sound_join_str if sound_only_flag else
                        markup_join_str if markup_only_flag else 
                        sound_markup_join_str if sound_markup_flag else ''))

            self.format_gray = (
                self.workbook.add_format({'bg_color': '#e6e6e6'}) if self.workbook else None)

        # If we need to export sounds and/or markups, we will need some sound/markup fields info.

        if sound_flag or markup_flag:

            self.storage = storage

            self.storage_f = (
                as_storage_file if __debug_flag__ else storage_file)

        self.ordering_type_id = (
            translation_gist_id_search('Ordering', self.session))

        if sound_flag:

            self.sound_type_id = (
                translation_gist_id_search('Sound', self.session))

        if markup_flag:

            self.markup_type_id = (
                translation_gist_id_search('Markup', self.session))

        self.is_order_field_dict = {}

        if __debug_flag__:

            # Simple text field ids query.

            self.text_field_query = session.query(
                Field.client_id, Field.object_id).filter_by(
                data_type_translation_gist_client_id = 1,
                data_type_translation_gist_object_id = 47,
                marked_for_deletion = False)

            # Simple list of text field ids.

            self.text_fields = self.text_field_query.all()

            # Temporary table with text field ids.

            self.text_field_id_table_name = (
                'text_field_id_table_' +
                str(uuid.uuid4()).replace('-', '_'))

            self.session.execute('''

                create temporary table {text_field_id_table_name} (
                  client_id BIGINT,
                  object_id BIGINT,
                  primary key (client_id, object_id))
                on commit drop;

                insert into {text_field_id_table_name}
                select client_id, object_id
                from field
                where
                data_type_translation_gist_client_id = 1 and
                data_type_translation_gist_object_id = 47 and
                marked_for_deletion = false;

                '''.format(
                    text_field_id_table_name = self.text_field_id_table_name))

    def is_order_field(self, field_id):
        """
        Checks if a field specified by its id is an ordering field.
        """

        result = self.is_order_field_dict.get(field_id)

        if result is not None:
            return result

        field = (

            Field.get(
                field_id,
                session = self.session))

        result = (
            field.data_type_translation_gist_id == self.ordering_type_id)

        self.is_order_field_dict[field_id] = result

        return result

    @staticmethod
    def fix_unicode(line):
        return line.encode('rtfunicode').decode('utf-8')

    def ready_worksheet(
        self,
        worksheet_name):
        """
        Prepares for exporting data to another worksheet.
        """

        self.worksheet = (
            self.workbook.add_worksheet(
                sanitize_worksheet_name(worksheet_name)))

        self.row = 0

    def ready_perspective(
        self,
        perspective,
        dictionary = None,
        worksheet_flag = True,
        list_flag = False,
        __debug_flag__ = False):
        """
        Prepares for saving data of lexical entries of another perspective.
        """

        # Initializing another worksheet, if required.

        if worksheet_flag:

            perspective_name = (
                perspective.get_translation(
                    self.locale_id, self.session))

            id_str = '_{0}_{1}'.format(
                perspective.client_id, perspective.object_id)

            sanitized_name = sanitize_worksheet_name(
                perspective_name, max_width = 31 - len(id_str))

            if self.workbook:
                self.worksheet = (
                    self.workbook.add_worksheet(
                        sanitize_worksheet_name(sanitized_name + id_str)))
            elif self.document:
                self.table = self.document.add_table(rows=1, cols=0, style="Table Grid")
            self.row = 0

        # Listing dictionary and perspective names, if required.

        if list_flag and self.workbook:

            if self.row > 0:
                self.row += 1

            if dictionary:

                self.worksheet.write(self.row, 0,
                    dictionary.get_translation(self.locale_id, self.session))

                self.row += 1

            self.worksheet.write(self.row, 0,
                perspective.get_translation(self.locale_id, self.session))

            self.row += 1

        # Getting field data.

        data_type_id_list = []

        if self.ordering_type_id:

            data_type_id_list.append(self.ordering_type_id)

        if (self.sound_flag and
            self.sound_type_id):

            data_type_id_list.append(self.sound_type_id)

        if (self.markup_flag and
            self.markup_type_id):

            data_type_id_list.append(self.markup_type_id)

        self.field_condition = (

            or_(
                DictionaryPerspectiveToField.field_id.in_(
                    sqlalchemy.text('select * from text_field_id_view')),

                *(
                    Field.data_type_translation_gist_id == data_type_id
                    for data_type_id in data_type_id_list)))

        field_query = (
                
            self.session

                .query(
                    DictionaryPerspectiveToField,
                    Field)

                .filter(
                    DictionaryPerspectiveToField.parent_id == perspective.id,
                    DictionaryPerspectiveToField.marked_for_deletion == False,
                    Field.id == DictionaryPerspectiveToField.field_id,
                    Field.marked_for_deletion == False,
                    self.field_condition)

                .order_by(DictionaryPerspectiveToField.position))

        if __debug_flag__:

            # Showing PostgreSQL's explain analyze, if required.

            row_list = (
                self.session.execute(
                    explain_analyze(field_query)).fetchall())

            log.debug(''.join(
                '\n' + row[0] for row in row_list))

        field_info_list = field_query.all()

        # All perspective fields we are to extract, and separately just the sound field ids.

        self.fields = [

            to_field 
            for to_field, field in field_info_list]

        if self.sound_flag:

            self.sound_field_id_set = set(

                field.id
                for to_field, field in field_info_list
                if field.data_type_translation_gist_id == self.sound_type_id)

        if self.markup_flag:

            self.markup_field_id_set = set(

                field.id
                for to_field, field in field_info_list
                if field.data_type_translation_gist_id == self.markup_type_id)

        # Etymology field, if required.

        self.etymology_field = (

            None if not self.cognates_flag else

            self.session.query(DictionaryPerspectiveToField).filter_by(
                parent_id=(perspective.client_id, perspective.object_id),
                marked_for_deletion=False,
                field_client_id=66,
                field_object_id=25).first())

        # To be sure, recursively ordering fields by subfield links, while preserving order.

        subfield_dict = collections.defaultdict(list)

        for field in self.fields:
            subfield_dict[field.self_id].append(field)

        field_set = set()
        field_list = []

        def f(field):

            if field in field_set:
                return

            field_set.add(field)
            field_list.append(field)

            for subfield in subfield_dict[field.id]:
                f(subfield)

        # NOTE: not just 'in subfield_dict[(None, None)]', as we request only a subset of fields and we can
        # get a subfield without a higher field, e.g. a markup field without its sound field.

        for field in self.fields:
            f(field)

        # Showing field ordering, if required.

        if __debug_flag__:

            field_info_list = [
                (field.field.get_translation(), field.id, field.self_id, field.field_id, field.position)
                for field in field_list]

            log.debug(
                '\nfield info:\n' +
                pprint.pformat(
                    field_info_list, width = 192))

        self.field_to_column = {
            field.field_id: counter
            for counter, field in enumerate(field_list)}

        # Making more space.

        if worksheet_flag and self.workbook:

            self.worksheet.set_column(
                0, len(self.fields) - 1, 13)

            if self.cognates_flag:

                self.worksheet.set_column(
                    len(self.fields), len(self.fields) + 4, 13)

        # Listing fields.

        column = 0
        tabs = list()
        cells = list()

        for field in self.fields:
            column_name = field.field.get_translation(self.locale_id, self.session)
            is_order = self.is_order_field(field.field_id)
            if self.workbook:
                self.worksheet.write(self.row, column, column_name)
                column += 1
            elif self.document:
                self.table.add_column(Inches(1) if is_order else Inches(3)).cells[0].text = column_name
            elif self.richtext:
                tabs.append(TabPropertySet.DEFAULT_WIDTH * (2 if is_order else 5))
                cells.append(Cell(Paragraph(self.fix_unicode(column_name))))

        if self.etymology_field:
            etymology_name = self.etymology_field.field.get_translation(self.locale_id, self.session)
            if self.workbook:
                self.worksheet.write(self.row, column, etymology_name)
            elif self.document:
                self.table.add_column(Inches(2)).cells[0].text = etymology_name
            elif self.richtext:
                tabs.append(TabPropertySet.DEFAULT_WIDTH * 5)
                cells.append(Cell(Paragraph(self.fix_unicode(etymology_name))))

        if self.richtext:
            self.table = Table(*tabs)
            self.table.AddRow(*cells)
            section = Section()
            section.append(self.table)
            self.richtext.Sections.append(section)

        self.row += 1

        self.perspective_fail_set = set()

        self.perspective_snd_fail_set = set()
        self.perspective_mrkp_fail_set = set()

        self.perspective_snd_mrkp_fail_set = set()

    def update_dictionary_info(
        self, perspective_id_set):
        """
        Ensures that we have required info for all given perspectives.
        """

        update_pid_set = set()

        for perspective_id in perspective_id_set:

            perspective_info = (
                self.perspective_info_dict.get(perspective_id))

            if perspective_info:

                order_key, name_str, xcript_fid, xlat_fid, thrd_fid, snd_fid, mrkp_fid = perspective_info

                if xcript_fid is None or xlat_fid is None:

                    self.perspective_fail_set.add(
                        (order_key, name_str))

                elif (
                    self.sound_flag and self.markup_flag and
                    (snd_fid is None or mrkp_fid is None)):

                    self.perspective_snd_mrkp_fail_set.add(
                        (order_key, name_str))

                elif (
                    self.sound_flag and
                    snd_fid is None):

                    self.perspective_snd_fail_set.add(
                        (order_key, name_str))

                elif (
                    self.markup_flag and
                    mrkp_fid is None):

                    self.perspective_mrkp_fail_set.add(
                        (order_key, name_str))

            else:

                update_pid_set.add(perspective_id)

        if not update_pid_set:
            return

        # Looking through all not already processed perspectives.

        query = (

            self.session

                .query(
                    Dictionary,
                    DictionaryPerspective)

                .filter(

                    tuple_(
                        DictionaryPerspective.client_id,
                        DictionaryPerspective.object_id)

                        .in_(update_pid_set),

                    DictionaryPerspective.marked_for_deletion == False,

                    Dictionary.id == DictionaryPerspective.parent_id,
                    Dictionary.marked_for_deletion == False))

        # For each perspective we compile ordering info and its title, with ordering based on standard
        # language order and dictionary/perspective creation times.

        insert_list = []

        for dictionary, perspective in query.all():

            order_index = (
                self.language_order_dict[dictionary.parent_id])

            order_key = (
                order_index, dictionary.created_at, perspective.created_at)

            perspective_name_str = (

                perspective.get_translation(
                    self.locale_id, self.session))

            name_str = (

                dictionary.get_translation(
                    self.locale_id, self.session) +

                ' › ' +
                perspective_name_str)

            # Getting info of perspective's fields.

            field_name_info_list = (

                self.session
                
                    .query(
                        Field.client_id,
                        Field.object_id,

                        func.jsonb_object_agg(
                            TranslationAtom.locale_id,
                            TranslationAtom.content)

                            .label('name'))

                    .filter(
                        DictionaryPerspectiveToField.parent_id == (perspective.client_id, perspective.object_id),
                        DictionaryPerspectiveToField.marked_for_deletion == False,
                        Field.id == DictionaryPerspectiveToField.field_id,
                        Field.marked_for_deletion == False,
                        self.field_condition,
                        Field.translation_gist_id == TranslationAtom.parent_id,
                        TranslationAtom.marked_for_deletion == False)

                    .group_by(
                        Field.client_id,
                        Field.object_id)

                    .order_by(
                        Field.client_id,
                        Field.object_id)

                    .all())

            text_field_position_list = (

                self.session

                    .query(
                        DictionaryPerspectiveToField.field_client_id,
                        DictionaryPerspectiveToField.field_object_id,
                        DictionaryPerspectiveToField.position)

                    .filter(
                        DictionaryPerspectiveToField.parent_id == (perspective.client_id, perspective.object_id),
                        DictionaryPerspectiveToField.marked_for_deletion == False,
                        DictionaryPerspectiveToField.field_id.in_(
                            sqlalchemy.text('select * from text_field_id_view')))

                    .group_by(
                        DictionaryPerspectiveToField.field_client_id,
                        DictionaryPerspectiveToField.field_object_id,
                        DictionaryPerspectiveToField.position)

                    .order_by(
                        DictionaryPerspectiveToField.position)

                    .all())

            log.debug(
                '\nfield_name_info_list:\n' +
                pprint.pformat(
                    field_name_info_list, width = 192))

            # Choosing perspective's transcription, translation, sound and markup fields.
            #
            # At first, we look for standard fields.

            transcription_fid = None
            translation_fid = None

            sound_fid = None
            markup_fid = None

            for field_cid, field_oid, name_dict in field_name_info_list:

                if field_cid == 66 and field_oid == 8:
                    transcription_fid = (field_cid, field_oid)

                elif field_cid == 66 and field_oid == 10:
                    translation_fid = (field_cid, field_oid)

                elif field_cid == 66 and field_oid == 12:
                    sound_fid = (field_cid, field_oid)

                elif field_cid == 66 and field_oid == 23:
                    markup_fid = (field_cid, field_oid)

            # If we don't have standard fields, we look through fields by names.

            if transcription_fid is None or translation_fid is None:

                en_xcript_fid = None
                en_xlat_fid = None

                en_snd_fid = None
                en_mrkp_fid = None

                ru_xcript_fid = None
                ru_xlat_fid = None

                ru_snd_fid = None
                ru_mrkp_fid = None

                for field_cid, field_oid, name_dict in field_name_info_list:

                    if '2' in name_dict:

                        en_name = (
                            re.sub(r'\W+', '', name_dict['2']).lower())
                        
                        if en_name == 'phonemictranscription':
                            en_xcript_fid = (field_cid, field_oid)

                        elif en_name == 'meaning':
                            en_xlat_fid = (field_cid, field_oid)

                        elif en_name == 'sound':
                            en_snd_fid = (field_cid, field_oid)

                        elif en_name == 'markup':
                            en_mrkup_fid = (field_cid, field_oid)

                    if '1' in name_dict:

                        ru_name = (
                            re.sub(r'\W+', '', name_dict['1']).lower())
                        
                        if ru_name == 'фонологическаятранскрипция':
                            ru_xcript_fid = (field_cid, field_oid)

                        elif ru_name == 'значение':
                            ru_xlat_fid = (field_cid, field_oid)

                        elif ru_name == 'звук':
                            ru_snd_fid = (field_cid, field_oid)

                        elif ru_name == 'разметка':
                            ru_mrkp_fid = (field_cid, field_oid)

                if en_xcript_fid is not None and en_xlat_fid is not None:

                    transcription_fid = en_xcript_fid
                    translation_fid = en_xlat_fid

                    sound_fid = en_snd_fid
                    markup_fid = en_mrkp_fid

                elif ru_xcript_fid is not None and ru_xlat_fid is not None:

                    transcription_fid = ru_xcript_fid
                    translation_fid = ru_xlat_fid

                    sound_fid = ru_snd_fid
                    markup_fid = ru_mrkp_fid

            # If we don't have fields with standard names, maybe this is paradigmatic perspective with
            # paradigmatic fields?

            if transcription_fid is None or translation_fid is None:

                en_xcript_fid = None
                en_xlat_fid = None

                en_snd_fid = None
                en_mrkp_fid = None

                ru_xcript_fid = None
                ru_xlat_fid = None

                ru_snd_fid = None
                ru_mrkp_fid = None

                for field_cid, field_oid, name_dict in field_name_info_list:

                    if '2' in name_dict:

                        en_name = (
                            re.sub(r'\W+', '', name_dict['2']).lower())
                        
                        if en_name == 'transcriptionofparadigmaticforms':
                            en_xcript_fid = (field_cid, field_oid)

                        elif en_name == 'translationofparadigmaticforms':
                            en_xlat_fid = (field_cid, field_oid)

                        elif en_name == 'soundofparadigmaticforms':
                            en_snd_fid = (field_cid, field_oid)

                        elif en_name == 'paradigmmarkup':
                            en_mrkp_fid = (field_cid, field_oid)

                    if '1' in name_dict:

                        ru_name = (
                            re.sub(r'\W+', '', name_dict['1']).lower())
                        
                        if ru_name == 'транскрипцияпарадигматическихформ':
                            ru_xcript_fid = (field_cid, field_oid)

                        elif ru_name == 'переводпарадигматическихформ':
                            ru_xlat_fid = (field_cid, field_oid)

                        elif ru_name == 'звукпарадигматическихформ':
                            ru_snd_fid = (field_cid, field_oid)

                        elif ru_name == 'разметкапарадигмы':
                            ru_mrkp_fid = (field_cid, field_oid)

                if en_xcript_fid is not None and en_xlat_fid is not None:

                    transcription_fid = en_xcript_fid
                    translation_fid = en_xlat_fid

                    sound_fid = en_snd_fid
                    markup_fid = en_mrkp_fid

                elif ru_xcript_fid is not None and ru_xlat_fid is not None:

                    transcription_fid = ru_xcript_fid
                    translation_fid = ru_xlat_fid

                    sound_fid = ru_snd_fid
                    markup_fid = ru_mrkp_fid

            # Finally, if it is a Starling-related perspective, we can try Starling-specific fields.
            #
            # No sound / markup fields for such a perspective.

            if ((transcription_fid is None or translation_fid is None) and
                perspective_name_str.lower().find('starling') != -1):

                en_xcript_fid = None
                en_xlat_fid = None

                ru_xcript_fid = None
                ru_xlat_fid = None

                for field_cid, field_oid, name_dict in field_name_info_list:

                    if '2' in name_dict:

                        en_name = (
                            re.sub(r'\W+', '', name_dict['2']).lower())
                        
                        if en_name == 'protoform':
                            en_xcript_fid = (field_cid, field_oid)

                        elif en_name == 'protoformmeaning':
                            en_xlat_fid = (field_cid, field_oid)

                    if '1' in name_dict:

                        ru_name = (
                            re.sub(r'\W+', '', name_dict['1']).lower())
                        
                        if ru_name == 'праформа':
                            ru_xcript_fid = (field_cid, field_oid)

                        elif ru_name == 'значениепраформы':
                            ru_xlat_fid = (field_cid, field_oid)

                if en_xcript_fid is not None and en_xlat_fid is not None:

                    transcription_fid = en_xcript_fid
                    translation_fid = en_xlat_fid

                elif ru_xcript_fid is not None and ru_xlat_fid is not None:

                    transcription_fid = ru_xcript_fid
                    translation_fid = ru_xlat_fid

            # Now trying to use regexp for transcription and translation fields

            if transcription_fid is None or translation_fid is None:

                en_xcript_fid = None
                en_xlat_fid = None

                ru_xcript_fid = None
                ru_xlat_fid = None

                def has_word(word, text):
                    return bool(re.search(r'\b' + word + r'\b', text))

                for field_cid, field_oid, name_dict in field_name_info_list:

                    if '2' in name_dict:

                        en_name = (
                            re.escape(name_dict['2']).lower())

                        if (has_word("transcription", en_name) or
                            has_word("word", en_name)):
                            en_xcript_fid = (field_cid, field_oid)

                        elif (has_word("translation", en_name) or
                              has_word("meaning", en_name)):
                            en_xlat_fid = (field_cid, field_oid)

                    if '1' in name_dict:

                        ru_name = (
                            re.escape(name_dict['1']).lower())

                        if (has_word("транскрипция", ru_name) or
                            has_word("слово", ru_name)):
                            ru_xcript_fid = (field_cid, field_oid)

                        elif (has_word("перевод", ru_name) or
                              has_word("значение", ru_name)):
                            ru_xlat_fid = (field_cid, field_oid)

                if en_xcript_fid is not None and en_xlat_fid is not None:

                    transcription_fid = en_xcript_fid
                    translation_fid = en_xlat_fid

                elif ru_xcript_fid is not None and ru_xlat_fid is not None:

                    transcription_fid = ru_xcript_fid
                    translation_fid = ru_xlat_fid

            # Failed to get transcription & translation fields, we will get first three fields.
            if transcription_fid is None or translation_fid is None:

                cognate_description_fids = [None, None, ('null', 'null')]
                for index, (field_cid, field_oid, _) in enumerate(text_field_position_list):
                    if index > 2:
                        break
                    cognate_description_fids[index] = (field_cid, field_oid)

                # Ok, failed to get three fields, we should remember it.
                if not all(cognate_description_fids):
                    self.perspective_fail_set.add(
                        (order_key, name_str))
            else:
                cognate_description_fids = [transcription_fid, translation_fid, ('null', 'null')]

            # Got transcription & translation fields, what about sound and/or markup?
            if all(cognate_description_fids):

                # If we are going to show both sound and markup fields for cognates, we require that the
                # markup field is a subfield of the sound field.

                if self.sound_flag and self.markup_flag:

                    if sound_fid is not None and markup_fid is not None:

                        Sound = aliased(DictionaryPerspectiveToField, name = 'Sound')
                        Markup = aliased(DictionaryPerspectiveToField, name = 'Markup')

                        sound_markup_check = (

                            self.session

                                .query(

                                    self.session

                                        .query(
                                            literal(1))

                                        .filter(
                                            Sound.parent_id == (perspective.client_id, perspective.object_id),
                                            Sound.marked_for_deletion == False,
                                            Sound.field_id == sound_fid,
                                            Markup.parent_id == (perspective.client_id, perspective.object_id),
                                            Markup.marked_for_deletion == False,
                                            Markup.field_id == markup_fid,
                                            Markup.self_id == Sound.id)

                                        .exists())

                                .scalar())

                        if not sound_markup_check:

                            sound_fid = None
                            markup_fid = None

                    if sound_fid is None or markup_fid is None:

                        self.perspective_snd_mrkp_fail_set.add(
                            (order_key, name_str))

                # Either a single sound or a single markup field.

                if (self.sound_flag and
                    sound_fid is None):

                    self.perspective_snd_fail_set.add(
                        (order_key, name_str))

                if (self.markup_flag and
                    markup_fid is None):

                    self.perspective_mrkp_fail_set.add(
                        (order_key, name_str))

                format_str = (
                    '({}, {}, {}, {}, {}, {}, {}, {}, {}, {}, {}, {})' if self.sound_flag and self.markup_flag else
                    '({}, {}, {}, {}, {}, {}, {}, {}, {}, {})' if self.sound_flag or self.markup_flag else
                    '({}, {}, {}, {}, {}, {}, {}, {})')

                format_arg_list = [perspective.client_id, perspective.object_id] + \
                                  [cognate_description_fids[i][j] for i in range(3) for j in range(2)]

                if self.sound_flag:

                    format_arg_list.extend([
                        sound_fid[0] if sound_fid is not None else 'null',
                        sound_fid[1] if sound_fid is not None else 'null'])

                if self.markup_flag:

                    format_arg_list.extend([
                        markup_fid[0] if markup_fid is not None else 'null',
                        markup_fid[1] if markup_fid is not None else 'null'])

                insert_list.append(
                    format_str.format(
                        *format_arg_list))

            self.perspective_info_dict[perspective.id] = (
                order_key, name_str,
                cognate_description_fids[0],
                cognate_description_fids[1],
                cognate_description_fids[2],
                sound_fid, markup_fid)

        # Updating perspectives' transcription / translation field info, if required.

        if insert_list:

            self.session.execute('''

                insert into
                {pid_fid_table_name}
                values
                {insert_str};

                '''.format(
                    pid_fid_table_name = self.pid_fid_table_name,
                    insert_str = ',\n'.join(insert_list)))

    def get_etymology_info(
        self,
        tag,
        published):
        """
        Gets info of etymologycally linked lexical entries, caches results.
        """

        # Updating perspectives' info.

        perspective_id_list = [

            (perspective_cid, perspective_oid)

            for perspective_cid, perspective_oid in (

                self.session
                
                    .execute(
                        self.sql_etymology_str_a, {
                            'tag': tag,
                            'publish': published})
                        
                    .fetchall())]

        self.update_dictionary_info(
            perspective_id_list)

        # Getting etymology group with transcriptions and translations.

        result_list = (

            self.session

                .execute(
                    self.sql_etymology_str_b.format(
                        '' if published is None else
                            ' and P.published = {}'.format(published),
                        '' if published is None else
                            ' and Ps.published = {}'.format(published),
                        '' if published is None else
                            ' and Pm.published = {}'.format(published)))

                .fetchall())

        # Compiling appropriately ordered cognate info.

        info_list = []

        for result_item in result_list:

            (
                entry_cid,
                entry_oid,
                perspective_cid,
                perspective_oid,
                transcription_list,
                translation_list,
                third_field_list) = (

                result_item[:7])

            sound_markup_list = result_item[-1]

            perspective_info = (
                    
                self.perspective_info_dict.get(
                    (perspective_cid, perspective_oid)))

            if perspective_info is None:
                continue

            order_key, name_str, _, _, _, _, _ = perspective_info

            # Getting ready to compose etymology info to inserted into XLSX.

            if transcription_list is None:
                transcription_list = []

            if translation_list is None:
                translation_list = []

            if third_field_list is None:
                third_field_list = []

            if sound_markup_list is None:
                sound_markup_list = []

            zip_arg_list = [

                [name_str],

                [t.strip()
                    for t in transcription_list],

                ['\'' + t.strip() + '\''
                    for t in translation_list],

                [t.strip()
                    for t in third_field_list]]

            # Processing sound and/or markup links.

            if self.sound_flag and self.markup_flag:

                sound_cell_list = []
                markup_cell_list = []

                for s_content, s_created_at, markup_content_list in sound_markup_list:

                    sound_cell_list.append(
                        [self.get_sound_link(s_content, s_created_at)])

                    markup_link_list = (

                        [[self.get_markup_link(*args)]
                            for args in markup_content_list
                            if args[0] is not None])

                    markup_cell_list.extend(markup_link_list)

                    if not markup_link_list:
                        markup_cell_list.append('')

                    elif len(markup_link_list) > 1:
                        sound_cell_list.extend([''] * (len(markup_link_list) - 1))

                zip_arg_list.extend([
                    sound_cell_list,
                    markup_cell_list])

            elif self.sound_flag:

                zip_arg_list.extend([

                    [[self.get_sound_link(*args)]
                        for args in sound_markup_list],

                    []])

            elif self.markup_flag:

                zip_arg_list.extend([

                    [[self.get_markup_link(*args)]
                        for args in sound_markup_list],

                    []])

            else:

                zip_arg_list.extend([[], []])

            # Preparing info of another cognate lexical entry.

            row_list = (
                    
                tuple(
                    itertools.zip_longest(
                        *zip_arg_list,
                        fillvalue = '')))

            info_list.append((
                order_key,
                (entry_cid, entry_oid),
                row_list))

        info_list.sort()

        log.debug(
            '\ninfo_list:\n' +
            pprint.pformat(
                info_list, width = 192))

        for _, entry_id, _ in info_list:
            self.etymology_dict[entry_id] = info_list

        return info_list

    def save_lexical_entry(
        self,
        entry,
        published = None,
        accepted = True,
        f_type = 'xlsx',
        __debug_flag__ = False):
        """
        Save data of a lexical entry of the current perspective.
        """

        rows_to_write = [[] for _ in self.fields]

        entities = (
            self.session.query(Entity).filter(
                Entity.parent_id == (entry.client_id, entry.object_id),
                Entity.marked_for_deletion == False))

        if published is not None or accepted is not None:

            entities = entities.join(PublishingEntity)

            if published is not None:
                entities = entities.filter(PublishingEntity.published == published)

            if accepted is not None:
                entities = entities.filter(PublishingEntity.accepted == accepted)

        # Arranging entities by subentity relation.

        entity_list = entities.all()
        subentity_dict = collections.defaultdict(list)

        for entity in entity_list:
            subentity_dict[entity.self_id].append(entity)

        def f(entity):
            """
            Recursively processes entities and subentities.
            """

            ent_field_ids = (
                entity.field_id)

            # Adding etymologycal info, if required.

            if ent_field_ids == (66, 25):

                if (self.etymology_field and
                    len(rows_to_write) == len(self.fields)):

                    entry_id = (entry.client_id, entry.object_id)

                    if entry_id in self.etymology_dict:

                        etymology_info = (
                            self.etymology_dict[entry_id])

                    else:

                        etymology_info = (

                            self.get_etymology_info(
                                entity.content, published))

                    # Dictionary name, transcriptions and translations of all linked entries.

                    start_index = (
                        len(rows_to_write))

                    rows_to_write.extend(
                        ([], [], [], [], [], []))

                    for _, cognate_entry_id, row_list in etymology_info:

                        if cognate_entry_id == entry_id:
                            continue

                        cell_format = None

                        if row_list and not row_list[0][1]:

                            cell_format = self.format_gray

                        for cell_list in row_list:

                            for index, value in enumerate(cell_list):
                                rows_to_write[start_index + index].append((value, cell_format))

                return {}

            # Text, sound / markup entity.

            row_count_dict = collections.Counter()

            column = (
                self.field_to_column.get(ent_field_ids))

            if column is not None:

                if (self.sound_flag and
                    entity.field_id in self.sound_field_id_set):

                    # Sound entity, getting sound file link, wrapping it in a list to show it's a link.

                    sound_link = (

                        self.get_sound_link(
                            entity.content,
                            entity.created_at))

                    rows_to_write[column].append([sound_link])

                elif (self.markup_flag and
                    entity.field_id in self.markup_field_id_set):

                    # Markup entity, getting markup file link, wrapping it in a list to show it's a link.

                    markup_link = (

                        self.get_markup_link(
                            entity.content,
                            entity.created_at))

                    rows_to_write[column].append([markup_link])

                else:

                    # Text entity, simply getting its text content.

                    rows_to_write[column].append(entity.content)

                row_count_dict[column] += 1

            # Now processing subentities, if we have any, and aligning column content if subentities'
            # data takes more than one row.

            for subentity in subentity_dict[entity.id]:

                for subcolumn, subcount in f(subentity).items():
                    row_count_dict[subcolumn] += subcount

            row_count = (
                max(row_count_dict.values(), default = 0))

            for column, count in row_count_dict.items():

                if count < row_count:

                    rows_to_write[column].extend(
                        [''] * (row_count - count))

                    row_count_dict[column] = row_count

            return row_count_dict

        # Recursively processing all entities of the entry, starting with top-level ones.

        for entity in subentity_dict[(None, None)]:
            f(entity)

        if __debug_flag__:

            log.debug(
                '\nrows_to_write:\n' +
                pprint.pformat(
                    rows_to_write, width = 192))

        def xlsx():
            # Writing out lexical entry data, if we have any.
            if any(rows_to_write):

                for cell_list in (
                    itertools.zip_longest(
                        *rows_to_write,
                        fillvalue = '')):

                    for index, cell in enumerate(cell_list):

                        # Tuple means with format.

                        if isinstance(cell, tuple):

                            value, cell_format = cell

                            if isinstance(value, list):

                                self.worksheet.write_url(
                                    self.row, index, './' + value[0], cell_format, value[0])

                            else:

                                self.worksheet.write(
                                    self.row, index, value, cell_format)

                        # List means a local link.

                        elif isinstance(cell, list):

                            self.worksheet.write_url(
                                self.row, index, './' + cell[0], None, cell[0])

                        else:

                            self.worksheet.write(
                                self.row, index, cell)

                    self.row += 1

        def docx():
            # Writing out lexical entry data, if we have any.
            if any(rows_to_write):

                for cell_list in (
                    itertools.zip_longest(
                        *rows_to_write,
                        fillvalue = '')):

                    row_cells = self.table.add_row().cells

                    for index, value in enumerate(cell_list):

                        # First checking if it's an order in case of unexpected yet possible irregular
                        # ordering entity values, like '' or None, which we still don't want to skip.

                        if (len(self.fields) > index and
                                self.is_order_field(self.fields[index].field_id)):

                            value = str(self.row)

                        else:

                            if isinstance(value, tuple):
                                value = value[0]

                            if not value or not isinstance(value, str):
                                continue

                        # Add columns if required
                        while len(row_cells) <= index:
                            self.table.add_column(Inches(2))

                        row_cells[index].text = value

                    self.row += 1

        def rtf():
            # Writing out lexical entry data, if we have any.
            if any(rows_to_write):
                frame = FramePropertySet(top=BorderPropertySet(width=20, style=BorderPropertySet.SINGLE))
                for cell_list in (
                    itertools.zip_longest(
                        *rows_to_write,
                        fillvalue = '')):

                    cells = list()
                    for index, value in enumerate(cell_list):
                        if (len(self.fields) > index and
                                self.is_order_field(self.fields[index].field_id)):

                            value = str(self.row)

                        else:

                            if isinstance(value, tuple):
                                value = value[0]

                            if not isinstance(value, str):
                                value = ''

                        cells.append(Cell(Paragraph(self.fix_unicode(value), frame)))

                    self.table.AddRow(*cells)
                    self.row += 1

        return (xlsx if f_type == 'xlsx' else
                docx if f_type == 'docx' else
                rtf)

    def get_zip_info(
        self,
        name,
        hash = None,
        date = None):
        """
        Prepares file for saving into the zip archive, changes file name if it is a duplicate.
        """

        zip_name = (

            self.zip_hash_dict.get(
                (name, hash)))

        if zip_name is not None:
            return zip_name

        if date is None:
            date = datetime.datetime.utcnow()

        name_count = (
            self.zip_name_dict[name])

        self.zip_name_dict[name] += 1

        if name_count >= 1:

            name_root, name_ext = (
                path.splitext(name))

            zip_name = (

                '{}_{}{}'.format(
                    name_root, name_count, name_ext))

        else:
            zip_name = name

        self.zip_hash_dict[(name, hash)] = zip_name

        zip_date = (
            date.year,
            date.month,
            date.day,
            date.hour,
            date.minute,
            date.second)

        return zipfile.ZipInfo(zip_name, zip_date)

    def get_sound_link(
        self,
        sound_url,
        created_at):
        """
        Processes linked sound file, adding it to the archive if necessary.
        """

        with self.storage_f(
            self.storage, sound_url) as sound_stream:

            sound_bytes = sound_stream.read()

        # Checking if we need to save the file, and if we need to rename it to avoid duplicate
        # names.

        zip_info = (
                
            self.get_zip_info(

                path.basename(
                    urllib.parse.urlparse(sound_url).path),

                hashlib.sha256(
                    sound_bytes).digest(),

                datetime.datetime.utcfromtimestamp(
                    created_at)))

        # Saving sound file to the archive, if required.

        if isinstance(zip_info, str):
            return zip_info

        if sndhdr.test_wav(
            sound_bytes, io.BytesIO(sound_bytes)):

            zip_info.compress_type = zipfile.ZIP_DEFLATED

        self.zip_file.writestr(zip_info, sound_bytes)

        return zip_info.filename

    def get_markup_link(
        self,
        markup_url,
        created_at):
        """
        Processes linked markup file, adding it to the archive if necessary.
        """

        with self.storage_f(
            self.storage, markup_url) as markup_stream:

            markup_bytes = markup_stream.read()

        # Checking if we need to save the file, and if we need to rename it to avoid duplicate
        # names.

        zip_info = (
                
            self.get_zip_info(

                path.basename(
                    urllib.parse.urlparse(markup_url).path),

                hashlib.sha256(
                    markup_bytes).digest(),

                datetime.datetime.utcfromtimestamp(
                    created_at)))

        # Saving markup file to the archive, if required.

        if isinstance(zip_info, str):
            return zip_info

        zip_info.compress_type = zipfile.ZIP_DEFLATED

        self.zip_file.writestr(zip_info, markup_bytes)

        return zip_info.filename

def write_xlsx(
    context,
    lex_dict,
    published,
    __debug_flag__ = False):

    for (_, lex) in sorted(lex_dict.items()):
        context.save_lexical_entry(
            lex,
            published,
            f_type = 'xlsx',
            __debug_flag__ = __debug_flag__)()

    # Writing out additional perspective info, if we have any.

    if context.perspective_fail_set:

        context.row += 1

        context.worksheet.write(
            context.row, 0,
            'Perspective{} without appropriate transcription / translation fields:'.format(
                '' if len(context.perspective_fail_set) == 1 else 's'))

        context.row += 1

        for _, name_str in sorted(context.perspective_fail_set):

            context.worksheet.write(
                context.row, 0, name_str)

            context.row += 1

    # The same for missing sound fields.

    if context.perspective_snd_fail_set:

        context.row += 1

        context.worksheet.write(
            context.row, 0,
            'Perspective{} without appropriate sound field:'.format(
                '' if len(context.perspective_snd_fail_set) == 1 else 's'))

        context.row += 1

        for _, name_str in sorted(context.perspective_snd_fail_set):

            context.worksheet.write(
                context.row, 0, name_str)

            context.row += 1

    # And for missing markup fields.

    if context.perspective_mrkp_fail_set:

        context.row += 1

        context.worksheet.write(
            context.row, 0,
            'Perspective{} without appropriate markup field:'.format(
                '' if len(context.perspective_mrkp_fail_set) == 1 else 's'))

        context.row += 1

        for _, name_str in sorted(context.perspective_mrkp_fail_set):

            context.worksheet.write(
                context.row, 0, name_str)

            context.row += 1

    # And for missing joint sound / markup field pairs.

    if context.perspective_snd_mrkp_fail_set:

        context.row += 1

        context.worksheet.write(
            context.row, 0,
            'Perspective{} without appropriate sound and markup fields:'.format(
                '' if len(context.perspective_snd_mrkp_fail_set) == 1 else 's'))

        context.row += 1

        for _, name_str in sorted(context.perspective_snd_mrkp_fail_set):

            context.worksheet.write(
                context.row, 0, name_str)

            context.row += 1


def write_docx(
    context,
    lex_dict,
    published,
    __debug_flag__ = False):

    for (_, lex) in sorted(lex_dict.items()):
        context.save_lexical_entry(
            lex,
            published,
            f_type = 'docx',
            __debug_flag__ = __debug_flag__)()


def write_rtf(
    context,
    lex_dict,
    published,
    __debug_flag__ = False):

    for (_, lex) in sorted(lex_dict.items()):
        context.save_lexical_entry(
            lex,
            published,
            f_type = 'rtf',
            __debug_flag__ = __debug_flag__)()


def compile_document(
    context,
    client_id,
    object_id,
    session,
    locale_id,
    published,
    storage = None,
    __debug_flag__ = False):
    """
    Compiles analysis results into xlsx/docx/rtf document.
    """

    dictionary = session.query(Dictionary).filter_by(id=(client_id, object_id),
                                                     marked_for_deletion=False).one()

    perspectives = session.query(DictionaryPerspective).filter_by(parent_id=(client_id, object_id),
                                                                  marked_for_deletion=False).all()

    # Processing all perspectives of the dictionary.

    for perspective in perspectives:

        context.ready_perspective(
            perspective,
            __debug_flag__ = __debug_flag__)

        lexical_entries = session.query(LexicalEntry, Entity).join(Entity).join(PublishingEntity) \
            .filter(LexicalEntry.parent_id == (perspective.client_id, perspective.object_id),
                    LexicalEntry.marked_for_deletion == False,
                    Entity.marked_for_deletion == False,
                    PublishingEntity.accepted == True)
        if published is not None:
            lexical_entries = lexical_entries.filter(PublishingEntity.published == published)

        lex_by_id = {}
        lex_by_order = {}

        for lex, entity in lexical_entries:
            lex_by_id[lex.id] = lex

            if context.is_order_field(entity.field_id):

                # Safeguarding against unexpected yet possible irregular null ordering entity values, which
                # should not appear but we can't guarantee against.

                lex_by_order[(entity.content or '', lex.id)] = lex

        lex_dict = lex_by_order if lex_by_order else lex_by_id

        if context.workbook:
            write_xlsx(context, lex_dict, published, __debug_flag__)
        elif context.document:
            write_docx(context, lex_dict, published, __debug_flag__)
            context.document.save(context.stream)
        elif context.richtext:
            write_rtf(context, lex_dict, published, __debug_flag__)
            # Write utf to bytes
            wrapper_file = codecs.getwriter('utf-8')(context.stream)
            context.richtext.write(wrapper_file)

    if context.workbook:
        context.workbook.close()


# @profile()
def save_dictionary(
    client_id,
    object_id,
    storage,
    sqlalchemy_url,
    task_key,
    cache_kwargs,
    dict_name,
    locale_id,
    published,
    sound_flag,
    markup_flag,
    f_type = 'xlsx',
    __debug_flag__ = False
):  # :(

    initialize_cache(cache_kwargs)
    task_status = TaskStatus.get_from_cache(task_key) if task_key else None

    engine = create_engine(sqlalchemy_url)
    register_after_fork(engine, engine.dispose)
    log = logging.getLogger(__name__)
    Session = sessionmaker(bind=engine)
    session = Session()

    if task_status:
        task_status.set(3, 20, 'Running async process')

    try:

        # Creating saving context, compiling dictionary data to a workbook.

        save_context = (
                
            Save_Context(
                locale_id,
                session,
                True,
                sound_flag,
                markup_flag,
                storage,
                f_type,
                __debug_flag__))

        if sound_flag or markup_flag:

            temporary_zip_file = (
                    
                tempfile.NamedTemporaryFile(
                    delete = False))

            save_context.zip_file = (
                zipfile.ZipFile(temporary_zip_file, 'w'))

            save_context.zip_name_dict = (
                collections.Counter())

            save_context.zip_hash_dict = {}

        compile_document(
            save_context,
            client_id,
            object_id,
            session,
            locale_id,
            published,
            storage,
            __debug_flag__ = __debug_flag__)

        # Name(s) of the resulting file(s) includes dictionary name, perspective name and current date.

        current_datetime = datetime.datetime.now(datetime.timezone.utc)

        result_filename = '{0} - {1:04d}.{2:02d}.{3:02d}'.format(
            dict_name[:64],
            current_datetime.year,
            current_datetime.month,
            current_datetime.day)

        table_filename = sanitize_filename(f"{result_filename}.{f_type}")

        if __debug_flag__:

            with open(table_filename, 'wb') as xlsx_file:

                save_context.stream.seek(0)
                copyfileobj(save_context.stream, xlsx_file)

        # Either adding XLSX file to the Zip archive...

        if sound_flag or markup_flag:

            zip_info = (

                save_context.get_zip_info(
                    table_filename))

            save_context.stream.seek(0)

            save_context.zip_file.writestr(
                zip_info,
                save_context.stream.read())

            save_context.zip_file.close()
            temporary_zip_file.close()

            zip_filename = (
                sanitize_filename(result_filename + '.zip'))

            zip_file_path = temporary_zip_file.name

            # Saving a local copy of the archive, if required.

            if __debug_flag__:

                shutil.move(
                    temporary_zip_file.name,
                    zip_filename)

                zip_file_path = zip_filename

            storage_temporary = storage['temporary']

            host = storage_temporary['host']
            bucket = storage_temporary['bucket']

            minio_client = (
                    
                minio.Minio(
                    host,
                    access_key = storage_temporary['access_key'],
                    secret_key = storage_temporary['secret_key'],
                    secure = True))

            object_name = (

                storage_temporary['prefix'] +
            
                '/'.join((
                    'save_dictionary',
                    '{:.6f}'.format(time.time()),
                    zip_filename)))

            (etag, version_id) = (

                minio_client.fput_object(
                    bucket,
                    object_name,
                    zip_file_path,
                    'application/zip'))

            url = (

                '/'.join((
                    'https:/',
                    host,
                    bucket,
                    object_name)))

            log.debug(
                '\nobject_name:\n{}'
                '\netag:\n{}'
                '\nversion_id:\n{}'
                '\nurl:\n{}'.format(
                    object_name,
                    etag,
                    version_id,
                    url))

            url_list = [url]

            if not __debug_flag__:

                os.remove(
                    temporary_zip_file.name)

        # Or saving it in the object storage.

        else:

            dictionary = session.query(Dictionary).filter_by(id=(client_id, object_id)).one()

            dict_status_atom = session.query(TranslationAtom).filter_by(
                parent_client_id=dictionary.state_translation_gist_client_id,
                parent_object_id=dictionary.state_translation_gist_object_id,
                locale_id=2).first()

            if not dict_status_atom:
                dict_status = 'translation_failure'
            else:
                dict_status = dict_status_atom.content

            if published is None:
                cur_folder = 'edit'
            elif published is True:
                cur_folder = 'view'
            else:
                cur_folder = 'should_be_impossible'

            cur_time = '{:.6f}'.format(time.time())

            storage_dir = (
                    
                path.join(
                    storage['path'],
                    'save_dictionary',
                    dict_status,
                    cur_folder,
                    cur_time))

            makedirs(storage_dir, exist_ok=True)

            # Storing file with the results.

            storage_path = path.join(storage_dir, table_filename)
            directory = path.dirname(storage_path)

            try:

                makedirs(directory)

            except OSError as exception:
                if exception.errno != EEXIST:
                    raise

            # If the name of the result file is too long, we try again with a shorter name.

            try:

                with open(storage_path, 'wb+') as workbook_file:

                    save_context.stream.seek(0)
                    copyfileobj(save_context.stream, workbook_file)

            except OSError as os_error:

                if os_error.errno != 36:
                    raise

                result_filename = '{0} - {1:04d}.{2:02d}.{3:02d}'.format(
                    dict_name[:32],
                    current_datetime.year,
                    current_datetime.month,
                    current_datetime.day)

                table_filename = sanitize_filename(f"{result_filename}.{f_type}")
                storage_path = path.join(storage_dir, table_filename)

                with open(storage_path, 'wb+') as workbook_file:

                    save_context.stream.seek(0)
                    copyfileobj(save_context.stream, workbook_file)

            # Successfully saved dictionary, finishing and returning links to files with results.

            url_list = [

                ''.join([
                    storage['prefix'],
                    storage['static_route'],
                    'save_dictionary', '/',
                    dict_status, '/',
                    cur_folder, '/',
                    cur_time, '/',
                    table_filename])]

        if task_status:
            task_status.set(4, 100, 'Finished', result_link_list=url_list)

    except Exception as exception:

        traceback_string = ''.join(traceback.format_exception(
            exception, exception, exception.__traceback__))[:-1]

        log.debug('save_dictionary: exception')
        log.debug(traceback_string)

        # If we failed to save the dictionary, we terminate with error.

        if task_status:

            task_status.set(4, 100,
                'Finished (ERROR), exception:\n' + traceback_string)

        return {'error': 'result compilation error'}

    session.commit()
    engine.dispose()


def get_json_tree(only_in_toc=True):

    cte_dict = get_cte_dict(only_in_toc)
    perspective_list = has_etymology(cte_dict['field_cte'])
    #json = compile_json(cte_dict)


# Getting cte for languages, dictionaries, perspectives and fields

def get_cte_dict(only_in_toc):

    # Getting root languages

    language_init = (
        SyncDBSession
            .query(
                Language,
                literal(0).label('level'))

            .filter(
                Language.parent_client_id == None,
                Language.parent_object_id == None,
                Language.marked_for_deletion == False)

            .cte(recursive=True))

    prnLanguage = aliased(language_init)
    subLanguage = aliased(Language)

    # Recursively getting tree of languages

    language_step = language_init.union_all(
        SyncDBSession
            .query(
                subLanguage,
                (prnLanguage.c.level + 1).label("level"))

            .filter(
                subLanguage.parent_client_id == prnLanguage.c.client_id,
                subLanguage.parent_object_id == prnLanguage.c.object_id,
                subLanguage.marked_for_deletion == False))

    if_only_in_toc = [language_step.c.additional_metadata['toc_mark'] == 'true'] if only_in_toc else []

    get_translation_atom = [
        TranslationGist.marked_for_deletion == False,
        TranslationAtom.parent_id == TranslationGist.id,
        func.length(TranslationAtom.content) > 0,
        TranslationAtom.marked_for_deletion == False ]

    language_сte = (
        SyncDBSession
            .query(
                func.min(language_step.c.level).label('language_level'),
                language_step.c.client_id.label('language_cid'),
                language_step.c.object_id.label('language_oid'),
                func.array_agg(TranslationAtom.content).label('language_title'))

            .filter(
                language_step.c.translation_gist_client_id == TranslationGist.client_id,
                language_step.c.translation_gist_object_id == TranslationGist.object_id,
                *get_translation_atom, *if_only_in_toc)

            .group_by(
                'language_cid',
                'language_oid')

            .cte())

    # Getting dictionaries with self titles

    dictionary_сte = (
        SyncDBSession
            .query(
                Dictionary.parent_client_id.label('language_cid'),
                Dictionary.parent_object_id.label('language_oid'),
                Dictionary.client_id.label('dictionary_cid'),
                Dictionary.object_id.label('dictionary_oid'),
                func.array_agg(TranslationAtom.content).label('dictionary_title'))

            .filter(
                Dictionary.parent_client_id == language_сte.c.language_cid,
                Dictionary.parent_object_id == language_сte.c.language_oid,
                Dictionary.marked_for_deletion == False,
                Dictionary.translation_gist_id == TranslationGist.id,
                *get_translation_atom)

            .group_by(
                'language_cid',
                'language_oid',
                'dictionary_cid',
                'dictionary_oid')

            .cte())

    # Getting perspectives with self titles

    perspective_сte = (
        SyncDBSession
            .query(
                DictionaryPerspective.parent_client_id.label('dictionary_cid'),
                DictionaryPerspective.parent_object_id.label('dictionary_oid'),
                DictionaryPerspective.client_id.label('perspective_cid'),
                DictionaryPerspective.object_id.label('perspective_oid'),
                func.array_agg(TranslationAtom.content).label('perspective_title'))

            .filter(
                DictionaryPerspective.parent_client_id == dictionary_сte.c.dictionary_cid,
                DictionaryPerspective.parent_object_id == dictionary_сte.c.dictionary_oid,
                DictionaryPerspective.marked_for_deletion == False,
                DictionaryPerspective.translation_gist_id == TranslationGist.id,
                *get_translation_atom)

            .group_by(
                'dictionary_cid',
                'dictionary_oid',
                'perspective_cid',
                'perspective_oid')

            .cte())

    # Getting fields with self title

    field_cte = (
        SyncDBSession
            .query(
                perspective_сte.c.perspective_cid,
                perspective_сte.c.perspective_oid,
                Field.client_id.label('field_cid'),
                Field.object_id.label('field_oid'),
                func.array_agg(func.lower(TranslationAtom.content)).label('field_title'),
                func.min(DictionaryPerspectiveToField.position).label('field_position'))

            .filter(
                DictionaryPerspectiveToField.parent_client_id == perspective_сte.c.perspective_cid,
                DictionaryPerspectiveToField.parent_object_id == perspective_сte.c.perspective_oid,
                DictionaryPerspectiveToField.marked_for_deletion == False,
                DictionaryPerspectiveToField.field_id == Field.id,
                Field.marked_for_deletion == False,
                Field.translation_gist_id == TranslationGist.id,
                *get_translation_atom, TranslationAtom.locale_id <= 2)

            .group_by(
                perspective_сte.c.perspective_cid,
                perspective_сte.c.perspective_oid,
                'field_cid', 'field_oid')

            .cte())

    return {
        'language_сte': language_сte,
        'dictionary_сte': dictionary_сte,
        'perspective_сte': perspective_сte,
        'field_cte': field_cte
    }


# Getting perspectives with transcription, translation and cognates

def has_etymology(field_cte):

    def has_word(word, text):
        return bool(re.search(r'\b' + word + r'\b', text))

    # Group fields by perspective
    fields_by_perspective = itertools.groupby(
        SyncDBSession.query(field_cte).yield_per(100),
        key=lambda x: (x[0], x[1]))

    for perspective_id, fields_group in fields_by_perspective:

        # Sorting fields by position
        fields_list = sorted(list(fields_group), key=lambda x: x[5])

        xcript_fid = None
        xlat_fid = None
        with_cognates = False

        for _, _, field_cid, field_oid, title, _ in fields_list:

            title = "; ".join(title)

            if xcript_fid is None and not has_word("affix", title):
                if (has_word("transcription", title) or
                        has_word("word", title) or
                        has_word("транскрипция", title) or
                        has_word("слово", title)):
                    xcript_fid = (field_cid, field_oid)
                    xcript_fname = title

            if xlat_fid is None and not has_word("affix", title):
                if (has_word("translation", title) or
                        has_word("meaning", title) or
                        has_word("перевод", title) or
                        has_word("значение", title)):
                    xlat_fid = (field_cid, field_oid)
                    xlat_fname = title

            if ((field_cid, field_oid) == (66, 25)):
                with_cognates = True

            if xcript_fid and xlat_fid and with_cognates:
                break

        if xcript_fid and xlat_fid and with_cognates:

            entities = (
                SyncDBSession
                    .query(
                        LexicalEntry.client_id,
                        LexicalEntry.object_id,
                        Entity.field_id,
                        Entity.content)

                    .filter(
                        LexicalEntry.parent_id == perspective_id,
                        Entity.parent_id == LexicalEntry.id,
                        Entity.field_id.in_([xcript_fid, xlat_fid]),
                        Entity.marked_for_deletion == False,
                        Entity.client_id == PublishingEntity.client_id,
                        Entity.object_id == PublishingEntity.object_id,
                        PublishingEntity.published == True,
                        PublishingEntity.accepted == True)

                    .yield_per(100))

            entities_by_lex = itertools.groupby(entities, key=lambda x: (x[0], x[1]))

            for lex_id, entities_group in entities_by_lex:

                linked_group = (
                    SyncDBSession
                        .execute(
                            f'select * from linked_group(66, 25, {lex_id[0]}, {lex_id[1]})')
                        .fetchall())

                entities_by_field = itertools.groupby(entities_group, key = lambda x: (x[2], x[3]))

                for field_id, group in entities_by_field:

                    field_text = [x[4] for x in group]

                    if field_id == xcript_fid:
                        xcript_text = field_text
                    elif field_id == xlat_fid:
                        xlat_text = field_text

                print(f"Perspective_id: {perspective_id}")
                print(f"{xcript_fname}: {xcript_text}")
                print(f"{xlat_fname}: {xlat_text}")
                print(f"Cognate_groups: {str(linked_group)}\n")
    """
    # Summary tree

    summary_сte = (
        SyncDBSession
            .query(
                language_сte.c.language_level,
                language_сte.c.language_cid,
                language_сte.c.language_oid,
                language_сte.c.language_title,

                dictionary_сte.c.dictionary_cid,
                dictionary_сte.c.dictionary_oid,
                dictionary_сte.c.dictionary_title,

                perspective_сte.c.perspective_cid,
                perspective_сte.c.perspective_oid,
                perspective_сte.c.perspective_title)

            .filter(
                language_сte.c.language_cid == dictionary_сte.c.language_cid,
                language_сte.c.language_oid == dictionary_сte.c.language_oid,
                dictionary_сte.c.dictionary_cid == perspective_сte.c.dictionary_cid,
                dictionary_сte.c.dictionary_oid == perspective_сte.c.dictionary_oid)

            .order_by(
                language_сte.c.language_level,
                language_сte.c.language_cid,
                language_сte.c.language_oid,

                dictionary_сte.c.dictionary_cid,
                dictionary_сte.c.dictionary_oid,

                perspective_сte.c.perspective_cid,
                perspective_сte.c.perspective_oid,)

            .yield_per(100))
    """