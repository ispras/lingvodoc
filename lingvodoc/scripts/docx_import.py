
# Standard library imports.

import ast
import getopt
import logging
import pprint
import re
import sys

# External imports.

import docx
import pympi
import pyramid.paster as paster

# Project imports.

from lingvodoc.models import (
    DBSession,
    Dictionary,
)


# Setting up logging, if we are not being run as a script.

if __name__ != '__main__':

    log = logging.getLogger(__name__)
    log.debug('module init')


def levenshtein(
    snippet_str,
    snippet_index,
    word_str,
    __debug_levenshtein_flag__ = False):
    """
    Matches word string to the snippet string via adjusted Levenshtein matching, with no penalties for
    snippet string skipping before and after match.
    """

    d = {(0, j): (j, 1e256)
        for j in range(len(word_str) + 1)}

    for i in range(len(snippet_str) - snippet_index):
        d[(i + 1, 0)] = (0, 1e256)

    minimum_distance = len(word_str)

    minimum_begin_index = 0
    minimum_end_index = 0

    for i in range(1, len(snippet_str) - snippet_index + 1):

        if __debug_levenshtein_flag__:

            log.debug(
                'd[{0}, 0]: {1}'.format(i, d[(i, 0)]))

        for j in range(1, len(word_str) + 1):

            # Matching current characters of the word and snippet strings.

            s_distance, s_begin_index = d[i - 1, j - 1]

            substitution_value = s_distance + (
                0 if snippet_str[snippet_index + i - 1] == word_str[j - 1] else 1)

            substitution_index = min(s_begin_index, i - 1)

            # Skipping current character from the snippet string.

            d_distance, d_begin_index = d[i - 1, j]

            deletion_value = d_distance + (
                1 if j < len(word_str) else 0)

            deletion_index = d_begin_index

            # Skipping current character from the word string.

            i_distance, i_begin_index = d[i, j - 1]

            insertion_value = i_distance + 1
            insertion_index = i_begin_index

            # Getting minimum.

            minimum_value = min(
                substitution_value,
                deletion_value,
                insertion_value)

            if minimum_value == deletion_value:

                operation_index = 1
                minimum_index = deletion_index

            elif minimum_value == insertion_value:

                operation_index = 2
                minimum_index = insertion_index

            else:

                operation_index = 0
                minimum_index = substitution_index

            d[(i, j)] = (minimum_value, minimum_index)

            # Showing edit distance computation details.

            if __debug_levenshtein_flag__:

                log.debug(
                    '\nd[{0}, {1}] (\'{18}\' & \'{14}\'): {4}'
                    '\n d[{5}, {6}] (\'{2}\' & \'{3}\'): {9} + {10}{11} (\'{12}\', \'{13}\')'
                    '\n d[{5}, {1}] (\'{2}\' & \'{14}\'): {15} + {16}{17}'
                    '\n d[{0}, {6}] (\'{18}\' & \'{3}\'): {19} + 1{20}'.format(
                    i, j,
                    snippet_str[snippet_index : snippet_index + i - 1] + '|' +
                        snippet_str[snippet_index + i - 1],
                    word_str[: j - 1] + '|' + word_str[j - 1],
                    d[(i, j)][0],
                    i - 1, j - 1,
                    snippet_str[snippet_index : snippet_index + i - 1] + '|',
                    word_str[: j - 1] + '|',
                    d[(i - 1, j - 1)][0],
                    0 if snippet_str[snippet_index + i - 1] == word_str[j - 1] else 1,
                    '*' if operation_index == 0 else '',
                    snippet_str[snippet_index + i - 1],
                    word_str[j - 1],
                    word_str[:j] + '|',
                    d[(i - 1, j)][0],
                    1 if j < len(word_str) else 0,
                    '*' if operation_index == 1 else '',
                    snippet_str[snippet_index : snippet_index + i] + '|',
                    d[(i, j - 1)][0],
                    '*' if operation_index == 2 else ''))

        # Checking if we have a new best matching.

        if d[i, len(word_str)][0] < minimum_distance:

            minimum_distance, minimum_begin_index = d[i, len(word_str)]
            minimum_end_index = i

        if minimum_distance == 0:
            break

    return (
        minimum_distance,
        minimum_begin_index,
        minimum_end_index)


def prepare_match_string(cell_str):
    """
    Processes string for matching, finding and marking portions in parentheses to be considered as
    optional during matching.
    """

    chr_list = []
    chr_index = 0

    for match in re.finditer(r'\([^()]*?\)', cell_str):

        for chr in re.sub(
            r'\W+', '', cell_str[chr_index : match.start()]):

            chr_list.append((chr, False))

        for chr in re.sub(
            r'\W+', '', match.group(0)):

            chr_list.append((chr, True))

        chr_index = match.end()

    for chr in re.sub(
        r'\W+', '', cell_str[chr_index:]):

        chr_list.append((chr, False))

    return chr_list


def format_match_string(marked_chr_list):
    """
    Formats list of marked characters as a string.
    """

    chr_list = []
    mark_prev = False

    for chr, mark in marked_chr_list:

        if mark != mark_prev:
            chr_list.append('(' if mark else ')')

        chr_list.append(chr)
        mark_prev = mark

    if mark_prev:
        chr_list.append(')')

    return ''.join(chr_list)


class State(object):
    """
    State of snippet table parsing.
    """

    def __init__(self, snippet_str, cell_list, row_index):
        """
        Initialization with the contents of the first snippet string.
        """

        self.snippet_count = 0
        self.snippet_chain = None

        self.snippet_str = snippet_str

        self.row_index = row_index
        self.row_list = [cell_list]

        self.d0 = []
        self.d1 = [0.999 * i for i in range(len(self.snippet_str) + 1)]

        self.word_list = []
        self.word_str = []

        self.total_value = 0
        self.snippet_value = 0

    def process_row(
        self,
        row_str,
        cell_list,
        row_index,
        no_parsing_flag,
        __debug_flag__ = False):
        """
        Processing another data string, splitting into a state when it's a word string and another state
        when it's a new snippet string.
        """

        # First, assuming that this data string is the next snippet string.

        if row_str:

            copy = State(row_str, cell_list, row_index)

            copy.snippet_chain = (
                (tuple(self.row_list), self.row_index),
                self.snippet_chain)

            copy.snippet_count = self.snippet_count + 1

            copy.total_value = self.total_value + self.d1[-1]

            yield copy

            if no_parsing_flag:
                return

        # Second, assuming that this data string is a word string.

        len_prev = len(self.word_str)

        self.word_list.append(row_str)
        self.word_str += row_str

        self.row_list.append(cell_list)

        # Updating Levenshtein alignment of snippet words to the snippet string.

        for i in range(len(row_str)):

            self.d0 = self.d1
            self.d1 = [len_prev + i + 1]

            for j in range(len(self.snippet_str)):

                # Matching current characters of the snippet string and the word string.

                s_cost = 0 if self.snippet_str[j][0] == row_str[i][0] else 1

                if s_cost and (self.snippet_str[j][1] or row_str[i][1]):
                    s_cost = 0.001

                s_value = self.d0[j] + s_cost

                # Skipping current character either from the snippet string or from the word string.

                d_value = self.d1[j] + (0.000999 if self.snippet_str[j][1] else 0.999)
                i_value = self.d0[j + 1] + (0.001 if row_str[i][1] else 1)

                self.d1.append(min(s_value, d_value, i_value))

                # Showing debug info, if required.

                if __debug_flag__:

                    log.debug((
                        format_match_string(self.snippet_str[:j]),
                        format_match_string(self.word_str[:len_prev + i]),
                        self.d0[j],
                        self.snippet_str[j][0],
                        row_str[i][0],
                        round(s_value, 6)))

                    log.debug((
                        format_match_string(self.snippet_str[:j]),
                        format_match_string(self.word_str[:len_prev + i + 1]),
                        self.d1[j],
                        self.snippet_str[j][0],
                        round(d_value, 6)))

                    log.debug((
                        format_match_string(self.snippet_str[:j + 1]),
                        format_match_string(self.word_str[:len_prev + i]),
                        self.d0[j + 1],
                        row_str[i][0],
                        round(i_value, 6)))

                    log.debug((
                        format_match_string(self.snippet_str[:j + 1]),
                        format_match_string(self.word_str[:len_prev + i + 1]),
                        round(min(s_value, d_value, i_value), 6)))

                    log.debug(self.d1)

        # Updating alignment value.

        if len(self.word_str) <= 0:
            self.snippet_value = 0

        elif len(self.word_str) > len(self.snippet_str):
            self.snippet_value = self.d1[-1]

        else:
            self.snippet_value = min(
                self.d1[len(self.word_str) : 2 * len(self.word_str)])

        yield self


def beam_search_step(
    state_list,
    cell_str,
    cell_list,
    row_index,
    beam_width,
    no_parsing_flag,
    __debug_beam_flag__ = False):
    """
    Another step of alignment beam search.
    """

    if not state_list:

        return [State(
            cell_str, cell_list, row_index)]

    # Sorting parsing states by the snippet they are parsing.

    state_dict = {}

    for state in state_list:

        for state_after in (

            state.process_row(
                cell_str,
                cell_list,
                row_index,
                no_parsing_flag)):

            index = state_after.row_index

            # Leaving only states with the best snippet histories.

            if (index not in state_dict or
                state_after.total_value < state_dict[index][0]):

                state_dict[index] = (state_after.total_value, [state_after])

            elif state_after.total_value == state_dict[index][0]:

                state_dict[index][1].append(state_after)

    state_list = []

    for value, state_after_list in state_dict.values():
        state_list.extend(state_after_list)

    # Showing snippet alignment beam search state, if required.

    if __debug_beam_flag__:

        log.debug('\n' + 
            pprint.pformat([(
                round(state.total_value + state.snippet_value, 6), 
                state.snippet_count,
                format_match_string(state.snippet_str),
                '|'.join(
                    format_match_string(word_str)
                    for word_str in state.word_list))
                for state in state_list],
                width = 384))

    # Leaving only a number of best states.

    state_list.sort(key = lambda state:
        (state.total_value + state.snippet_value, state.snippet_count))

    return state_list[:beam_width]


def parse_table(
    row_list,
    limit = None,
    no_header_flag = False,
    no_parsing_flag = False,
    __debug_beam_flag__ = False):
    """
    Tries to parse snippet data represented as a table.
    """

    # Removing any snippet alignment marks, if we have any.
 
    for cell_list in row_list:
        for i in range(len(cell_list)):
 
            match = re.match(r'\(__\d+__\)\s*', cell_list[i])
     
            if match:
                cell_list[i] = cell_list[i][match.end():]

    state_list = []
    beam_width = 32

    # Going through snippet data.

    row_sequence = (
        enumerate(row_list) if no_header_flag else
        enumerate(row_list[1:], 1))

    for row_index, cell_list in row_sequence:

        if limit and row_index > limit:
            break

        if not any(cell_list[:3]):
            continue

        cell_str = (
            prepare_match_string(
                cell_list[0].lower()))

        # Updating alignment search on another row.

        state_list = (

            beam_search_step(
                state_list,
                cell_str,
                cell_list,
                row_index,
                beam_width,
                no_parsing_flag,
                __debug_beam_flag__))

    # Returning final parsing search state.

    return state_list


def parse_by_paragraphs(
    row_list,
    limit = None,
    __debug_flag__ = False,
    __debug_beam_flag__ = False):
    """
    Tries to parse snippet data with paragraph separation inside table cells.
    """

    # Splitting row texts by paragraphs.

    line_row_list = []
    line_row_count = 0

    for cell_list in row_list[1:]:

        if limit and line_row_count >= limit:
            break

        paragraph_list_list = [
            re.split(r'[^\S\n]*\n\s*', text)
            for text in cell_list]

        how_many = max(
            len(paragraph_list)
            for paragraph_list in paragraph_list_list[:3])

        # Iterating over aligned paragraphs in adjacent cells.

        line_rank_list = []

        for i in range(how_many):

            line_cell_list = []

            for paragraph_list in paragraph_list_list:

                if i < len(paragraph_list):

                    # Removing snippet alignment mark, if there is one present.

                    cell_str = paragraph_list[i]
                    match = re.match(r'\(__\d+__\)\s*', cell_str)

                    line_cell_list.append(
                        cell_str[match.end():] if match else
                        cell_str)

                else:
                    line_cell_list.append('')

            # Another line row, if it is non-empty.

            if any(line_cell_list):

                line_rank_list.append(line_cell_list)
                line_row_count += 1

                if limit and line_row_count >= limit:
                    break

        line_row_list.append(line_rank_list)

    # Showing what we have, if required.

    if __debug_flag__:

        log.debug(
            '\nrow_list:\n{0}'.format(
            pprint.pformat(
                row_list, width = 196)))

    state_list = []
    beam_width = 32

    line_row_count = 0

    # Going through snippet data.

    for row_index, line_rank_list in enumerate(line_row_list):

        if limit and line_row_count >= limit:
            break

        for line_index, line_cell_list in enumerate(line_rank_list):

            line_cell_str = (
                prepare_match_string(
                    line_cell_list[0].lower()))

            # Updating alignment search on another row.

            state_list = (
                beam_search_step(
                    state_list,
                    line_cell_str,
                    line_cell_list,
                    (row_index, line_index),
                    beam_width,
                    __debug_beam_flag__))

    # Returning final parsing search state.

    return state_list


class Docx2EafError(Exception):

    def __init__(self, message):
        super().__init__(message)


def docx2eaf(
    docx_path,
    eaf_file_path,
    separate_by_paragraphs_flag = False,
    modify_docx_flag = False,
    all_tables_flag = False,
    no_header_flag = False,
    no_parsing_flag = False,
    check_file_path = None,
    check_docx_file_path = None,
    limit = None,
    __debug_flag__ = False,
    __debug_beam_flag__ = False,
    __debug_eaf_flag__ = False):
    """
    Converts .docx file of the right structure to a 5-tier .eaf corpus file.
    """

    log.debug(
        '\ndocx_path: {0}'.format(docx_path))

    try:
        document = docx.Document(docx_path)

    except docx.opc.exceptions.PackageNotFoundError:
        raise Docx2EafError('input file is not a .docx format file')

    if len(document.tables) <= 0:
        raise Docx2EafError('.docx file does not have any tables')

    # Accessing info of the first table, or all tables, depending on the options.
    #
    # Counting only unique cells because apparently some .docx documents can have repeating cells in their
    # structure.

    row_list = []

    table_list = (
            
        document.tables if all_tables_flag else
        document.tables[:1])

    for table_index, table in enumerate(table_list):

        column_count = len(set(table.rows[0].cells))
        row_count = len(set(table.columns[0].cells))

        table_cell_list = list(table._cells)

        source_cell_list = []
        source_cell_set = set()
        
        for cell in table_cell_list:

            if cell not in source_cell_set:

                source_cell_list.append(cell)
                source_cell_set.add(cell)

        # Checking for non-uniform rows / columns.

        if len(source_cell_list) != column_count * row_count:

            error_str = (

                '\nTable ({0}): rows and / or columns are uneven, '
                '{1} rows, {2} columns, {3} != {1} * {2} cells.'.format(
                    table_index,
                    row_count,
                    column_count,
                    len(source_cell_list)))

            log.error(error_str)

            raise Docx2EafError(error_str)

        row_list.extend(

            [cell.text
                for cell in source_cell_list[
                    i * column_count : (i + 1) * column_count]]

                for i in range(row_count))

        log.debug(
            '\ntable ({}): {} columns, {} rows, {} cells'.format(
                table_index,
                column_count,
                row_count,
                len(source_cell_list)))

    # Processing this info.

    if not no_header_flag:

        header_list = row_list[0]

        log.debug(
            '\nheader: {0}'.format(header_list))

    if separate_by_paragraphs_flag:

        state_list = (

            parse_by_paragraphs(
                row_list,
                limit,
                __debug_flag__,
                __debug_beam_flag__))

    else:

        state_list = (

            parse_table(
                row_list,
                limit,
                no_header_flag,
                no_parsing_flag,
                __debug_beam_flag__))

    # Showing final alignment search state, if required.

    if __debug_beam_flag__:

        log.debug('\n' + 
            pprint.pformat([(
                round(state.total_value + state.snippet_value, 6), 
                state.snippet_count,
                format_match_string(state.snippet_str),
                '|'.join(
                    format_match_string(word_str)
                    for word_str in state.word_list))
                for state in state_list],
                width = 384))

    # Getting all parsed snippets, if we need them.

    if (eaf_file_path is not None or
        check_file_path is not None or
        check_docx_file_path is not None or
        modify_docx_flag):

        if not state_list:

            log.debug('\nno data')
            return

        best_state = state_list[0]

        snippet_chain = (
            (tuple(best_state.row_list), best_state.row_index),
            best_state.snippet_chain)

        snippet_list = []

        # Compiling snippet list, showing it, if required.

        while snippet_chain is not None:

            (row_tuple, row_index), snippet_chain = snippet_chain
            snippet_list.append((list(row_tuple), row_index))

        snippet_list.reverse()

        if __debug_flag__:

            log.debug(
                '\nsnippet_list:\n{0}'.format(
                pprint.pformat(
                    snippet_list, width = 196)))

    # Saving parsed alignment, if required.

    if check_file_path is not None:

        with open(
            check_file_path, 'w', encoding = 'utf-8') as check_file:

            check_file.write('\n')

            # Showing each parsed snippet.

            for i, (snippet_value_list, snippet_value_index) in enumerate(snippet_list):

                check_file.write(
                    '{0}\n'.format(i + 1))

                value = snippet_value_list[0]

                check_file.write(
                    (value if isinstance(value, str) else value[0]) + '\n')
                
                for value in snippet_value_list[1:]:

                    check_file.write('  ' +
                        (value if isinstance(value, str) else value[0]) + '\n')

                check_file.write('\n')

    # Saving parsing alignment as Docx file, if required.

    if (check_docx_file_path is not None and
        not separate_by_paragraphs_flag and
        not all_tables_flag):

        check_docx = docx.Document()

        check_table = check_docx.add_table(
            rows = row_count - 1 + len(snippet_list),
            cols = 3)

        table_cell_list = check_table._cells
        table_cell_index = 0

        # Exporting all parsed snippets with their numbers.

        for i, (snippet_row_list, snippet_row_index) in enumerate(snippet_list):

            table_cell_list[table_cell_index].text = '{0}'.format(i + 1)
            table_cell_index += 3

            for cell_list in snippet_row_list:

                for table_cell, snippet_cell in zip(
                    table_cell_list[table_cell_index : table_cell_index + 3],
                    cell_list):

                    table_cell.text = snippet_cell

                table_cell_index += 3

        check_docx.save(check_docx_file_path)

    # Saving parsed snippets as the standard 5-tier EAF structure.

    if eaf_file_path is not None:

        log.debug('\n' + pprint.pformat(snippet_list, width = 196))

        eaf = pympi.Elan.Eaf()

        eaf.add_linguistic_type('text_top_level')
        eaf.add_linguistic_type('symbolic_association', 'Symbolic_Association', False)
        eaf.add_linguistic_type('word_translation_included_in', 'Included_In')

        eaf.remove_linguistic_type('default-lt')

        # Showing linguistic types info, if required.

        if __debug_eaf_flag__:

            log.debug(
                '\nget_linguistic_type_names(): {0}'.format(eaf.get_linguistic_type_names()))

            log.debug(''.join(
                '\nget_parameters_for_linguistic_type({0}): {1}'.format(
                    repr(name),
                    eaf.get_parameters_for_linguistic_type(name))
                for name in eaf.get_linguistic_type_names()))

        eaf.add_tier('text', 'text_top_level')
        eaf.add_tier('other text', 'symbolic_association', 'text')
        eaf.add_tier('literary translation', 'symbolic_association', 'text')
        eaf.add_tier('translation', 'word_translation_included_in', 'text')
        eaf.add_tier('transcription', 'symbolic_association', 'translation')
        eaf.add_tier('word', 'symbolic_association', 'translation')

        eaf.remove_tier('default')

        # Showing tier info, if required.

        if __debug_eaf_flag__:

            log.debug(
                '\nget_tier_names(): {0}'.format(eaf.get_tier_names()))

            log.debug(''.join(
                '\nget_parameters_for_tier({0}): {1}'.format(
                    repr(name),
                    eaf.get_parameters_for_tier(name))
                for name in eaf.get_tier_names()))

        # Compiling annotation data.

        step = 75
        position = step

        for snippet_value_list, snippet_value_index in snippet_list:

            # Snippet base texts.

            text, text_other, text_translation = snippet_value_list[0]
            duration = len(text) * step

            eaf.add_annotation(
                'text', position, position + duration, text)

            eaf.add_ref_annotation(
                'other text', 'text', position, text_other)

            eaf.add_ref_annotation(
                'literary translation', 'text', position, text_translation)

            # Preparing to create annotations for snippet words.

            translation_position = position

            translation_length = (

                sum(len(text_list[0] or text_list[2] or text_list[1])
                    for text_list in snippet_value_list[1:]) +

                len(snippet_value_list) - 2)

            translation_position = position
            translation_step = duration // translation_length

            # Snippet words.

            for text_list in snippet_value_list[1:]:

                word, word_other, translation = text_list

                translation_duration = (
                        
                    round(
                        max(len(word or translation or word_other), 1) *
                        translation_step))

                eaf.add_annotation(
                    'translation',
                    translation_position,
                    translation_position + translation_duration,
                    translation)

                eaf.add_ref_annotation(
                    'transcription', 'translation', translation_position, word_other)

                eaf.add_ref_annotation(
                    'word', 'translation', translation_position, word)

                translation_position += (
                    translation_duration + translation_step)

            # Ready to go to the next snippet.

            position += duration + step

        # Showing annotation info, if required.

        if __debug_eaf_flag__:

            log.debug(''.join(
                '\nget_annotation_data_for_tier({0}):\n{1}'.format(
                    repr(name),
                    eaf.get_annotation_data_for_tier(name)[:4])
                for name in eaf.get_tier_names()))

        eaf.header['TIME_UNITS'] = 'milliseconds'

        eaf.to_file(eaf_file_path)

    # Modifying source Docx file with alignment marks, if required.

    if modify_docx_flag:

        if not separate_by_paragraphs_flag:

            for i, (snippet_row_list, snippet_row_index) in enumerate(snippet_list):

                mark_str = '(__{0}__)\n'.format(i + 1)
                cell_index = snippet_row_index * column_count

                for j, cell in enumerate(
                    source_cell_list[cell_index : cell_index + 3]):

                    # Right now can't do something like
                    #
                    #   cell.paragraphs[0].insert_paragraph_before(mark_str),
                    #
                    # because, if there is a mark there already, we should delete it, and tracking this
                    # deletion across all possible paragraphs and runs in the cell is too high complexity.

                    cell.text = mark_str + snippet_row_list[0][j]

            document.save(docx_path)

        # When tables are separated by paragraphs.

        else:

            snippet_index = 0
            snippet_row_index, snippet_rank_index = snippet_list[snippet_index][1]

            for row_index, cell_list in enumerate(row_list[1:]):

                # Along the lines of data extraction from such tables, see 'parse_by_paragraphs()' function.

                paragraph_list_list = [
                    re.split(r'([^\S\n]*\n\s*)', text)
                    for text in cell_list]

                for i, paragraph_list in enumerate(paragraph_list_list):

                    paragraph_list.append('')

                    paragraph_list_list[i] = list(
                        zip(paragraph_list[::2], paragraph_list[1::2]))

                line_list_list = [[]
                    for text in cell_list]

                how_many = max(
                    len(paragraph_list)
                    for paragraph_list in paragraph_list_list[:3])

                # Iterating over aligned paragraphs in adjacent cells.

                line_rank_count = 0

                for i in range(how_many):

                    line_cell_list = []

                    if (snippet_index is not None and
                        row_index == snippet_row_index and
                        line_rank_count == snippet_rank_index):

                        mark_str = '(__{0}__)\n'.format(snippet_index + 1)

                        for line_list in line_list_list:
                            line_list.append(mark_str)

                        # Next snippet.

                        snippet_index += 1

                        if snippet_index >= len(snippet_list):
                            snippet_index = None

                        else:
                            snippet_row_index, snippet_rank_index = (
                                snippet_list[snippet_index][1])

                    for paragraph_list, line_list in zip(
                        paragraph_list_list, line_list_list):

                        if i < len(paragraph_list):

                            # Removing previous snippet alignment mark, if there is one present.

                            cell_str, separator_str = paragraph_list[i]
                            match = re.match(r'\(__\d+__\)\s*', cell_str)

                            if match:
                                cell_str = cell_str[match.end():]

                            line_cell_list.append(cell_str)
                            line_list.append(cell_str + separator_str)

                    # Another line row, if it is non-empty.

                    if any(line_cell_list):
                        line_rank_count += 1

                    else:
                        for line_list in line_list_list:
                            line_list.pop()

                # Replacing contents of another table cell.

                cell_index = (row_index + 1) * column_count

                for cell, line_list in zip(
                    source_cell_list[cell_index : cell_index + 3],
                    line_list_list):

                    match = re.fullmatch(
                        r'(.*?)[^\S\n]*\n[^\S\n]*', line_list[0], re.DOTALL)

                    cell.text = match.group(1) if match else line_list[0]

                    # Splitting text into distinct paragraphs because otherwise at least LibreOffice writer
                    # starts to take too much time to process resulting documents.

                    for line in line_list[1:]:

                        match = re.fullmatch(
                            r'(.*?)[^\S\n]*\n[^\S\n]*', line, re.DOTALL)

                        cell.add_paragraph(
                            match.group(1) if match else line, 'Normal')

            # Saving Docx file updates.

            document.save(docx_path)


def main_import(args):
    """
    Test import of 5-tier data from a Docx file.
    """

    opt_list, arg_list = (
        getopt.gnu_getopt(args, '', [
            'all-tables',
            'check-docx-file=',
            'check-file=',
            'debug',
            'debug-beam',
            'debug-eaf',
            'eaf-file=',
            'limit=',
            'modify-docx-file',
            'no-db',
            'no-header',
            'no-parsing',
            'separate-by-paragraphs']))

    opt_dict = dict(opt_list)

    # Parsing command-line options.

    docx_path = arg_list[0]

    check_file_path = opt_dict.get('--check-file')
    check_docx_file_path = opt_dict.get('--check-docx-file')
    eaf_file_path = opt_dict.get('--eaf-file')

    limit = (
        ast.literal_eval(opt_dict['--limit'])
            if '--limit' in opt_dict else None)

    modify_docx_flag = '--modify-docx-file' in opt_dict
    separate_by_paragraphs_flag = '--separate-by-paragraphs' in opt_dict

    all_tables_flag = '--all-tables' in opt_dict
    no_header_flag = '--no-header' in opt_dict
    no_parsing_flag = '--no-parsing' in opt_dict

    __debug_flag__ = '--debug' in opt_dict
    __debug_beam_flag__ = '--debug-beam' in opt_dict
    __debug_eaf_flag__ = '--debug-eaf' in opt_dict

    # Processing specified Docx file.

    docx2eaf(
        docx_path,
        eaf_file_path,
        separate_by_paragraphs_flag,
        modify_docx_flag,
        all_tables_flag,
        no_header_flag,
        no_parsing_flag,
        check_file_path,
        check_docx_file_path,
        limit,
        __debug_flag__,
        __debug_beam_flag__,
        __debug_eaf_flag__)


def main_eaf(args):
    """
    Showing structure of a specified Eaf file.
    """

    for eaf_path in args:

        log.debug(
            '\neaf_path: {0}'.format(eaf_path))

        eaf = pympi.Elan.Eaf(eaf_path)

        log.debug(
            '\nget_controlled_vocabulary_names(): {0}'.format(eaf.get_controlled_vocabulary_names()))

        log.debug(
            '\nget_external_ref_names(): {0}'.format(eaf.get_external_ref_names()))

        log.debug(
            '\nget_languages(): {0}'.format(eaf.get_languages()))

        log.debug(
            '\nget_lexicon_ref_names(): {0}'.format(eaf.get_lexicon_ref_names()))

        log.debug(
            '\nget_licenses(): {0}'.format(eaf.get_licenses()))

        log.debug(
            '\nget_linguistic_type_names(): {0}'.format(eaf.get_linguistic_type_names()))

        log.debug(
            '\nget_linked_files(): {0}'.format(eaf.get_linked_files()))

        log.debug(
            '\nget_locales(): {0}'.format(eaf.get_locales()))

        log.debug(
            '\nget_properties(): {0}'.format(eaf.get_properties()))

        log.debug(
            '\nget_secondary_linked_files(): {0}'.format(eaf.get_secondary_linked_files()))

        log.debug(
            '\nget_tier_names(): {0}'.format(eaf.get_tier_names()))

        log.debug('\n' +
            pprint.pformat(eaf.linguistic_types, width = 196))

        # L-type and tier parameters.

        log.debug(''.join(
            '\nget_parameters_for_linguistic_type({0}): {1}'.format(
                repr(name),
                eaf.get_parameters_for_linguistic_type(name))
            for name in eaf.get_linguistic_type_names()))

        log.debug(''.join(
            '\nget_tier_ids_for_linguistic_type({0}): {1}'.format(
                repr(name),
                eaf.get_tier_ids_for_linguistic_type(name))
            for name in eaf.get_linguistic_type_names()))

        log.debug(''.join(
            '\nget_parameters_for_tier({0}): {1}'.format(
                repr(name),
                eaf.get_parameters_for_tier(name))
            for name in eaf.get_tier_names()))

        # Select annotations.

        log.debug(''.join(
            '\nget_annotation_data_for_tier({0}):\n{1}'.format(
                repr(name),
                eaf.get_annotation_data_for_tier(name)[:4])
            for name in eaf.get_tier_names()))

        # Average time interval per character.

        total_duration = 0
        total_length = 0

        for name in eaf.get_tier_names():

            tier_duration = 0
            tier_length = 0

            for annotation in eaf.get_annotation_data_for_tier(name):
                begin, end, text = annotation[:3]

                tier_duration += end - begin
                tier_length += len(text)

            log.debug(
                '\ntier {0}: {1:.3f} / {2} -> {3:.3f}'.format(
                repr(name),
                tier_duration / 1000.0,
                tier_length,
                tier_duration / (tier_length * 1000)))

            total_duration += tier_duration
            total_length += tier_length

        log.debug(
            '\ntotal: {0:.3f} / {1} -> {2:.3f}'.format(
            total_duration / 1000.0,
            total_length,
            total_duration / (total_length * 1000)))


# If we are being run as a script.

if __name__ == '__main__':

    if (len(sys.argv) > 1 and
        sys.argv[1] == '-config'):

        # We have a configuration file; initializing DB, if required, and logging.

        config_path = sys.argv[2]

        if sys.argv[3] != '-no-db':

            pyramid_env = paster.bootstrap(config_path)
            arg_list = sys.argv[3:]

        else:
            arg_list = sys.argv[4:]

        paster.setup_logging(config_path)
        log = logging.getLogger(__name__)

    else:

        # No config file, so just logging to stdout.

        arg_list = sys.argv[1:]

        log_root = logging.getLogger()
        log_root.setLevel(logging.DEBUG)

        log_handler = logging.StreamHandler(sys.stdout)
        log_handler.setLevel(logging.DEBUG)

        log_formatter = (
                
            logging.Formatter(
                '%(asctime)s %(levelname)-5.5s [%(name)s][%(threadName)s] '
                '%(pathname)s:%(lineno)d: %(message)s'))

        log_handler.setFormatter(log_formatter)
        log_root.addHandler(log_handler)

        log = logging.getLogger(__name__)

    # Doing what we need.

    if len(arg_list) <= 0:

        log.info(
            '\nPlease specify a command to execute.')

    elif arg_list[0] == 'import':

        main_import(arg_list[1:])

    elif arg_list[0] == 'eaf':

        main_eaf(arg_list[1:])

    else:

        log.warning(
            '\nUnknown command \'{0}\'.'.format(arg_list[0]))

